#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
独立的 ITU 干扰分析脚本
无需依赖其他配置文件，可直接运行

使用方法:
    python standalone_itu_analyzer.py /path/to/image.png

输出:
    - interference_report.md (与输入图片同目录)
    - interference_report.docx (与输入图片同目录)
"""

import os
import sys
import asyncio
import json
import re
import shutil
import subprocess
from datetime import datetime
from typing import Dict, List, Any, Optional
from dataclasses import dataclass, asdict
from enum import Enum

from PIL import Image
from autogen_core import Image as AutoGenImage
from autogen_agentchat.agents import AssistantAgent
from autogen_agentchat.conditions import TextMentionTermination
from autogen_agentchat.teams import RoundRobinGroupChat
from autogen_agentchat.messages import MultiModalMessage
from autogen_ext.models.openai import OpenAIChatCompletionClient

# ============================================================================
# API 配置
# ============================================================================
LLM_API_KEY = "sk-ant-api03-rHPBzv2u8cIlfGpYDpzefxjl1pT9GASmlKqBdt-sCAEsByKWPP-OFQBvVJdgs6ANqJuvs96aoh-Xz3nlejfE_A"
LLM_BASE_URL = "https://api.aicodemirror.com/api/gemini"
LLM_MODEL_NAME = "gemini-2.5-flash"

# 代理配置（如不需要可注释掉）
HTTP_PROXY = "http://10.192.54.148:7897"

# ============================================================================
# Agent 角色定义
# ============================================================================
class AgentRole(Enum):
    """Agent角色定义"""
    PARSER = "parser"
    ANALYSIS = "analysis"
    REVIEW = "review"
    REPORT = "report"


@dataclass
class AgentConfig:
    """Agent配置"""
    name: str
    role: AgentRole
    description: str
    max_tokens: int
    temperature: float
    system_message: str
    input_source: Optional[str] = None
    output_format: str = "text"
    termination_marker: str = ""


# ============================================================================
# 数据流管理
# ============================================================================
class DataFlowManager:
    """管理agent间的数据流转"""

    def __init__(self):
        self.data_flow = {}
        self.audit_log = []

    def record_input(self, agent_name: str, input_data: Any, source: str) -> None:
        """记录agent的输入"""
        if agent_name not in self.data_flow:
            self.data_flow[agent_name] = {}
        self.data_flow[agent_name]["input"] = {
            "data": input_data,
            "source": source,
            "timestamp": datetime.now().isoformat(),
        }

    def record_output(self, agent_name: str, output_data: Any, format_type: str = "text") -> None:
        """记录agent的输出"""
        if agent_name not in self.data_flow:
            self.data_flow[agent_name] = {}
        self.data_flow[agent_name]["output"] = {
            "data": output_data,
            "format": format_type,
            "timestamp": datetime.now().isoformat(),
        }

    def get_data_flow_summary(self) -> Dict[str, Any]:
        """获取数据流转摘要"""
        summary = {}
        for agent_name, flow in self.data_flow.items():
            summary[agent_name] = {
                "has_input": "input" in flow,
                "has_output": "output" in flow,
            }
        return summary

    def export_audit_log(self, filepath: str) -> None:
        """导出审计日志"""
        with open(filepath, "w", encoding="utf-8") as f:
            json.dump(self.audit_log, f, indent=2, ensure_ascii=False)

    def export_data_flow(self, filepath: str) -> None:
        """导出数据流转记录"""
        serializable_flow = {}
        for agent_name, flow in self.data_flow.items():
            serializable_flow[agent_name] = {
                "input": {
                    "source": flow.get("input", {}).get("source"),
                    "timestamp": flow.get("input", {}).get("timestamp"),
                    "data_preview": str(flow.get("input", {}).get("data"))[:200],
                } if "input" in flow else None,
                "output": {
                    "format": flow.get("output", {}).get("format"),
                    "timestamp": flow.get("output", {}).get("timestamp"),
                    "data_preview": str(flow.get("output", {}).get("data"))[:200],
                } if "output" in flow else None,
            }

        with open(filepath, "w", encoding="utf-8") as f:
            json.dump(serializable_flow, f, indent=2, ensure_ascii=False)

    @staticmethod
    def _get_timestamp() -> str:
        """获取当前时间戳"""
        return datetime.now().isoformat()


# ============================================================================
# Agent 配置构建器
# ============================================================================
class AgentConfigBuilder:
    """Agent配置构建器"""

    @staticmethod
    def build_parser_config() -> AgentConfig:
        return AgentConfig(
            name="parser_agent",
            role=AgentRole.PARSER,
            description="Parse ITU interference map and extract structured data",
            max_tokens=1024,
            temperature=0.2,
            system_message=(
                "You are the Parser Agent. Your ONLY job is to analyze the interference map image and extract structured data.\n\n"
                "CRITICAL: Output ONLY valid JSON (no markdown, no explanations). The JSON must contain:\n"
                "{\n"
                '  "basic_info": {\n'
                '    "report_date": "YYYY-MM-DD",\n'
                '    "time_range": "HH:MM-HH:MM or description",\n'
                '    "system_type": "Ground Terminal System or Satellite System",\n'
                '    "monitoring_metric": "CINR/CIR/CNR/EPFD/PFD/INR/etc",\n'
                '    "filename": "extracted from image or context"\n'
                "  },\n"
                '  "numerical_data": {\n'
                '    "metric_name": "extracted metric",\n'
                '    "min_value": number,\n'
                '    "max_value": number,\n'
                '    "avg_value": number,\n'
                '    "range": number,\n'
                '    "unit": "dB or other"\n'
                "  },\n"
                '  "temporal_characteristics": {\n'
                '    "abnormal_periods": ["period1", "period2"],\n'
                '    "duration_type": "short-lived/intermittent/sustained",\n'
                '    "magnitude": "low/moderate/high"\n'
                "  },\n"
                '  "visual_patterns": [\n'
                '    "pattern1 with specifics",\n'
                '    "pattern2 with specifics"\n'
                "  ]\n"
                "}\n\n"
                "INSTRUCTIONS:\n"
                "1. Look for COLORBAR on the right side - read MIN and MAX values\n"
                "2. Calculate AVG = (Min + Max) / 2\n"
                "3. Calculate RANGE = Max - Min\n"
                "4. Identify temporal anomalies from the heatmap\n"
                "5. List visual patterns observed\n\n"
                "Output ONLY the JSON, nothing else. End with: [PARSER_DONE]"
            ),
            input_source="image",
            output_format="json",
            termination_marker="[PARSER_DONE]",
        )

    @staticmethod
    def build_analysis_config() -> AgentConfig:
        return AgentConfig(
            name="analysis_agent",
            role=AgentRole.ANALYSIS,
            description="Analyze parsed data and identify patterns and interference",
            max_tokens=1536,
            temperature=0.3,
            system_message=(
                "You are the Analysis Agent. Your job is to analyze the structured data from Parser Agent.\n\n"
                "INPUT: You will receive JSON data from Parser Agent containing numerical data and visual patterns.\n"
                "OUTPUT: Generate analysis in this format:\n\n"
                "## Data Analysis Results\n\n"
                "**Interference Presence**: [Yes/No/Uncertain] - [reason based on data]\n\n"
                "### Numerical Analysis\n"
                "- Min Value: [value] dB\n"
                "- Max Value: [value] dB\n"
                "- Average: [value] dB\n"
                "- Range: [value] dB\n"
                "- Assessment: [Are values typical? Any concerning patterns?]\n\n"
                "### Temporal Characteristics\n"
                "- Abnormal Periods: [specific times/regions]\n"
                "- Duration: [short-lived/intermittent/sustained]\n"
                "- Magnitude: [low/moderate/high]\n\n"
                "### Pattern Analysis\n"
                "[2-3 paragraphs analyzing visual patterns, variability, and potential sources]\n\n"
                "### Potential Interference Sources\n"
                "1. [hypothesis with evidence from figure]\n"
                "2. [hypothesis with evidence]\n"
                "3. [hypothesis with evidence]\n\n"
                "IMPORTANT: Do NOT re-analyze the image. Use ONLY the data provided by Parser Agent.\n"
                "End with: [ANALYSIS_DONE]"
            ),
            input_source="parser",
            output_format="text",
            termination_marker="[ANALYSIS_DONE]",
        )

    @staticmethod
    def build_review_config() -> AgentConfig:
        return AgentConfig(
            name="review_agent",
            role=AgentRole.REVIEW,
            description="Review analysis for ITU compliance and logical consistency",
            max_tokens=1024,
            temperature=0.3,
            system_message=(
                "You are the Review Agent. Your job is to verify the analysis from Analysis Agent.\n\n"
                "INPUT: You will receive analysis results from Analysis Agent.\n"
                "OUTPUT: Generate review in this format:\n\n"
                "## Compliance Review\n\n"
                "### ITU Standard Compliance\n"
                "- **Compliance Status**: [Compliant/Non-Compliant/Partially Compliant/Marginal]\n"
                "- **Key Findings**: [summary of compliance assessment]\n"
                "- **Risk Level**: [High/Moderate/Low]\n\n"
                "[1 paragraph explaining compliance assessment based on ITU standards]\n\n"
                "### Evidence Summary\n"
                "- [Visual cue 1 with specifics]\n"
                "- [Visual cue 2 with specifics]\n"
                "- [Visual cue 3 with specifics]\n\n"
                "### Logical Consistency Check\n"
                "[Assessment of whether analysis is logically consistent and well-supported]\n\n"
                "### Recommendations\n"
                "1. **[Action 1]**: [Purpose and expected impact]\n"
                "2. **[Action 2]**: [Purpose and expected impact]\n"
                "3. **[Action 3]**: [Purpose and expected impact]\n\n"
                "End with: [REVIEW_DONE]"
            ),
            input_source="analysis",
            output_format="text",
            termination_marker="[REVIEW_DONE]",
        )

    @staticmethod
    def build_report_config() -> AgentConfig:
        return AgentConfig(
            name="report_agent",
            role=AgentRole.REPORT,
            description="Generate final formatted report",
            max_tokens=1536,
            temperature=0.2,
            system_message=(
                "You are the Report Agent. Your job is to generate the final formatted report.\n\n"
                "INPUT: You will receive review results from Review Agent.\n"
                "OUTPUT: Generate the complete report in markdown format:\n\n"
                "# Interference Analysis Report\n\n"
                "## 1. Basic Information\n"
                "| Field | Value |\n"
                "|-------|-------|\n"
                "| Report date | [date] |\n"
                "| System type | [type] |\n"
                "| Monitoring metric | [metric] |\n"
                "| Research institution | Institute of Space Internet, Fudan University |\n\n"
                "## 2. Data Analysis\n"
                "[Include analysis results from Review Agent]\n\n"
                "## 3. Compliance Assessment\n"
                "[Include compliance review from Review Agent]\n\n"
                "## 4. Conclusions and Recommendations\n"
                "[Include recommendations from Review Agent]\n\n"
                "IMPORTANT: Integrate all previous agent outputs into a cohesive, well-formatted report.\n"
                "Use **bold** for key values and statuses.\n"
                "End with: [REPORT_DONE]"
            ),
            input_source="review",
            output_format="markdown",
            termination_marker="[REPORT_DONE]",
        )

    @staticmethod
    def build_all_configs() -> Dict[str, AgentConfig]:
        return {
            "parser": AgentConfigBuilder.build_parser_config(),
            "analysis": AgentConfigBuilder.build_analysis_config(),
            "review": AgentConfigBuilder.build_review_config(),
            "report": AgentConfigBuilder.build_report_config(),
        }

    @staticmethod
    def export_configs(filepath: str) -> None:
        """导出所有配置到JSON文件"""
        configs = AgentConfigBuilder.build_all_configs()
        serializable_configs = {}
        for name, config in configs.items():
            config_dict = asdict(config)
            config_dict["role"] = config_dict["role"].value
            serializable_configs[name] = config_dict

        with open(filepath, "w", encoding="utf-8") as f:
            json.dump(serializable_configs, f, indent=2, ensure_ascii=False)


# ============================================================================
# Agent 序列验证器
# ============================================================================
class AgentSequenceValidator:
    """验证agent序列的有效性"""

    @staticmethod
    def validate_sequence(configs: List[AgentConfig]) -> bool:
        """验证agent序列"""
        if not configs:
            print("[ERROR] Agent sequence is empty")
            return False

        if configs[0].role != AgentRole.PARSER:
            print("[ERROR] First agent must be Parser")
            return False

        if configs[-1].role != AgentRole.REPORT:
            print("[ERROR] Last agent must be Report")
            return False

        for i in range(1, len(configs)):
            prev_agent = configs[i - 1]
            curr_agent = configs[i]
            expected_source = prev_agent.role.value
            if curr_agent.input_source != expected_source:
                print(
                    f"[ERROR] Agent {curr_agent.name} expects input from {curr_agent.input_source}, "
                    f"but previous agent {prev_agent.name} is {expected_source}"
                )
                return False

        print("[INFO] Agent sequence validation passed")
        return True

    @staticmethod
    def print_sequence_info(configs: List[AgentConfig]) -> None:
        """打印agent序列信息"""
        print("\n=== Agent Sequence Information ===")
        for i, config in enumerate(configs):
            print(f"\n[{i + 1}] {config.name}")
            print(f"    Role: {config.role.value}")
            print(f"    Input Source: {config.input_source}")
            print(f"    Output Format: {config.output_format}")
            print(f"    Max Tokens: {config.max_tokens}")
            print(f"    Temperature: {config.temperature}")
            print(f"    Termination Marker: {config.termination_marker}")


# ============================================================================
# 辅助函数
# ============================================================================
def configure_proxies() -> None:
    """配置 HTTP 代理"""
    os.environ["HTTP_PROXY"] = HTTP_PROXY
    os.environ["HTTPS_PROXY"] = HTTP_PROXY
    print(f"[INFO] HTTP(S) proxy configured: {HTTP_PROXY}")


def parse_image_info(filename: str) -> Dict:
    """从文件名解析图像信息"""
    parts = filename.split("_")
    constellation = parts[0] if parts else "Unknown"
    terminal_type = "Ground Terminal System" if "earth" in filename.lower() else "Satellite System"

    analysis_types_map = {
        "cinr": "CINR",
        "cir": "CIR",
        "cnr": "CNR",
        "epfd": "EPFD",
        "inr": "INR",
        "link_count": "Link Count",
        "pfd": "PFD",
        "temp": "Delta T/T",
    }
    analysis_type = next(
        (v for k, v in analysis_types_map.items() if k in filename.lower()), "Unknown"
    )

    return {
        "constellation": constellation,
        "terminal_type": terminal_type,
        "analysis_type": analysis_type,
        "filename": filename,
    }


def compress_image(image_path: str, max_size: int = 800, quality: int = 80) -> str:
    """压缩图像以减少 token 消耗"""
    img = Image.open(image_path)
    width, height = img.size
    if max(width, height) > max_size:
        scale = max_size / max(width, height)
        new_width = int(width * scale)
        new_height = int(height * scale)
        img = img.resize((new_width, new_height), Image.Resampling.LANCZOS)
        print(f"[INFO] Image resized: {width}x{height} -> {new_width}x{new_height}")

    compressed_path = image_path.replace(".png", "_compressed.jpg")
    img.convert("RGB").save(compressed_path, "JPEG", quality=quality, optimize=True)

    original_size = os.path.getsize(image_path) / 1024
    compressed_size = os.path.getsize(compressed_path) / 1024
    print(
        f"[INFO] Image compressed: {original_size:.1f}KB -> {compressed_size:.1f}KB (saved {100*(1-compressed_size/original_size):.1f}%)"
    )

    return compressed_path


def create_model_client(max_tokens: int = 2048, temperature: float = 0.3) -> OpenAIChatCompletionClient:
    """创建 LLM 客户端"""
    base_url = LLM_BASE_URL.rstrip("/")
    if base_url.endswith("/v1"):
        base_url = base_url[:-3]

    return OpenAIChatCompletionClient(
        model=LLM_MODEL_NAME,
        base_url=base_url,
        api_key=LLM_API_KEY,
        model_capabilities={
            "vision": True,
            "function_calling": True,
            "json_output": True,
        },
        create_args={
            "max_tokens": max_tokens,
            "temperature": temperature,
        },
    )


# ============================================================================
# 主管道类
# ============================================================================
class StandaloneAnalysisPipeline:
    """独立的干扰分析管道"""

    def __init__(self, image_info: dict, current_date: str, output_dir: str):
        self.image_info = image_info
        self.current_date = current_date
        self.output_dir = output_dir
        self.data_flow_manager = DataFlowManager()
        self.agents = {}
        self.agent_configs = {}

    def _create_agents_from_configs(self) -> Dict[str, AssistantAgent]:
        """从配置创建 agents"""
        configs = AgentConfigBuilder.build_all_configs()
        agents = {}

        for name, config in configs.items():
            client = create_model_client(
                max_tokens=config.max_tokens,
                temperature=config.temperature,
            )
            agent = AssistantAgent(
                config.name,
                model_client=client,
                description=config.description,
                system_message=config.system_message,
            )
            agents[name] = agent
            self.agent_configs[name] = config

        return agents

    async def run(self, image_path: str) -> str:
        """运行分析管道"""
        print(f"[INFO] Loading image: {image_path}")
        pil_image = Image.open(image_path)
        autogen_image = AutoGenImage(pil_image)

        # 创建 agents
        self.agents = self._create_agents_from_configs()

        # 验证 agent 序列
        configs_list = list(self.agent_configs.values())
        if not AgentSequenceValidator.validate_sequence(configs_list):
            print("[ERROR] Agent sequence validation failed")
            return ""

        AgentSequenceValidator.print_sequence_info(configs_list)

        # 记录初始输入
        self.data_flow_manager.record_input(
            "parser_agent",
            f"Image: {os.path.basename(image_path)}",
            "image",
        )

        # 构建初始提示
        initial_prompt = f"""Context Information:
- Constellation: {self.image_info['constellation']}
- Terminal Type: {self.image_info['terminal_type']}
- Analysis Type: {self.image_info['analysis_type']}
- Filename: {self.image_info['filename']}
- Date: {self.current_date}

IMPORTANT: This is a heat map / time-series plot showing interference metric values.
Look carefully at the COLORBAR (legend) on the side of the figure to identify min/max values.

Please analyze this interference map and extract structured data."""

        # 创建 GroupChat
        print("[INFO] Starting multi-agent pipeline...")
        termination = TextMentionTermination(text="[REPORT_DONE]")

        agent_sequence = [
            self.agents["parser"],
            self.agents["analysis"],
            self.agents["review"],
            self.agents["report"],
        ]

        group_chat = RoundRobinGroupChat(
            agent_sequence,
            termination_condition=termination,
        )

        # 启动对话
        task = MultiModalMessage(content=[initial_prompt, autogen_image], source="user")

        try:
            print("[INFO] Parser Agent analyzing image...")
            result = await group_chat.run(task=task)
            messages = result.messages

            print(f"[DEBUG] Total messages: {len(messages)}")

            # 提取最终报告
            final_report = self._extract_final_report(messages)

            if not final_report:
                print("[WARNING] No report generated")
                return ""

            print(f"[INFO] Report generated: {len(final_report)} chars")

            # 保存 Markdown 报告
            md_path = os.path.join(self.output_dir, "interference_report.md")
            with open(md_path, "w", encoding="utf-8") as f:
                f.write(final_report)
            print(f"[INFO] Markdown saved: {md_path}")

            # 转换为 Word 文档
            docx_path = os.path.join(self.output_dir, "interference_report.docx")
            self._create_docx_with_image(md_path, docx_path, image_path)
            print(f"[INFO] DOCX saved: {docx_path}")

            print("==== Report generation finished ====")
            return final_report

        except Exception as e:
            print(f"[ERROR] Pipeline failed: {e}")
            import traceback
            traceback.print_exc()
            return ""

    def _extract_final_report(self, messages: List) -> str:
        """提取并聚合最终报告"""
        agent_outputs = {
            "parser_agent": None,
            "analysis_agent": None,
            "review_agent": None,
            "report_agent": None,
        }

        for msg in messages:
            if isinstance(msg.content, str) and msg.content.strip():
                source = getattr(msg, "source", None)
                if not source:
                    source = getattr(msg, "agent_name", None)
                if not source:
                    source = getattr(msg, "name", None)

                if source in agent_outputs:
                    content = msg.content
                    for marker in ["[PARSER_DONE]", "[ANALYSIS_DONE]", "[REVIEW_DONE]", "[REPORT_DONE]"]:
                        content = content.replace(marker, "")
                    content = content.strip()
                    if content:
                        agent_outputs[source] = content

        parser_output = agent_outputs.get("parser_agent")
        analysis_output = agent_outputs.get("analysis_agent")
        review_output = agent_outputs.get("review_agent")
        report_output = agent_outputs.get("report_agent")

        return self._aggregate_report(parser_output, analysis_output, review_output, report_output)

    def _parse_text_to_table_data(self, text: str) -> Dict[str, str]:
        """
        从原始文本中提取键值对数据。

        Args:
            text: 原始文本内容

        Returns:
            提取的键值对字典
        """
        table_data = {}

        # 尝试多种格式的键值对提取
        patterns = [
            r'^\\s*([A-Za-z_\\s]+):\\s*(.+)$',  # Key: Value
            r'^\\s*([A-Za-z_\\s]+)\\s*=\\s*(.+)$',  # Key = Value
            r'^\\s*-\\s*([A-Za-z_\\s]+):\\s*(.+)$',  # - Key: Value
            r'^\\s*\\*\\s*([A-Za-z_\\s]+):\\s*(.+)$',  # * Key: Value
        ]

        lines = text.split('\\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue

            for pattern in patterns:
                match = re.match(pattern, line, re.IGNORECASE)
                if match:
                    key = match.group(1).strip().title()
                    value = match.group(2).strip()
                    value = value.strip('`"\\''')
                    if key and value:
                        table_data[key] = value
                    break

        # 如果没有提取到数据，尝试从 JSON 代码块中提取
        if not table_data:
            json_match = re.search(r'```(?:json)?\\s*(\\{[\\s\\S]*?\\})\\s*```', text)
            if json_match:
                try:
                    json_data = json.loads(json_match.group(1))
                    table_data = self._flatten_json(json_data)
                except (json.JSONDecodeError, ValueError):
                    pass

        return table_data

    def _flatten_json(self, data: dict, prefix: str = "") -> Dict[str, str]:
        """
        将嵌套的 JSON 数据扁平化为键值对。

        Args:
            data: JSON 数据
            prefix: 键前缀

        Returns:
            扁平化的键值对字典
        """
        result = {}
        for key, value in data.items():
            full_key = f"{prefix}{key}".replace('_', ' ').title() if prefix else key.replace('_', ' ').title()
            if isinstance(value, dict):
                result.update(self._flatten_json(value, f"{full_key} - "))
            elif isinstance(value, list):
                result[full_key] = ', '.join(str(v) for v in value)
            else:
                result[full_key] = str(value)
        return result

    def _add_structured_data_section_v2(self, report_parts: List[str], parser_data: dict,
                                        section_num: int, subsection: int) -> None:
        """添加结构化数据章节（从 JSON 解析），带小标题编号"""
        # Numerical Data Table
        if "numerical_data" in parser_data:
            report_parts.append(f"### {section_num}.{subsection} Numerical Measurements")
            report_parts.append("")
            report_parts.append("**Summary:** The table below presents the key statistical parameters extracted from the monitoring data.")
            report_parts.append("")
            num_data = parser_data["numerical_data"]
            report_parts.append("| Parameter | Value | Unit | Description |")
            report_parts.append("|-----------|-------|------|-------------|")
            report_parts.append(f"| Metric Type | {num_data.get('metric_name', 'N/A')} | - | Primary interference metric |")
            report_parts.append(f"| Minimum | {num_data.get('min_value', 'N/A')} | {num_data.get('unit', 'dB')} | Lowest recorded value |")
            report_parts.append(f"| Maximum | {num_data.get('max_value', 'N/A')} | {num_data.get('unit', 'dB')} | Highest recorded value |")
            report_parts.append(f"| Average | {num_data.get('avg_value', 'N/A')} | {num_data.get('unit', 'dB')} | Mean value over monitoring period |")
            report_parts.append(f"| Range | {num_data.get('range', 'N/A')} | {num_data.get('unit', 'dB')} | Difference between max and min |")
            report_parts.append("")
            subsection += 1

        # Temporal Characteristics
        if "temporal_characteristics" in parser_data:
            report_parts.append(f"### {section_num}.{subsection} Temporal Characteristics")
            report_parts.append("")
            report_parts.append("**Summary:** The temporal analysis identifies patterns and anomalies in the time-series data.")
            report_parts.append("")
            temporal = parser_data["temporal_characteristics"]
            abnormal = ', '.join(temporal.get('abnormal_periods', [])) if temporal.get("abnormal_periods") else "None identified"
            report_parts.append("| Characteristic | Value | Interpretation |")
            report_parts.append("|----------------|-------|----------------|")
            report_parts.append(f"| Abnormal Periods | {abnormal} | Time periods with significant deviations |")
            report_parts.append(f"| Duration Pattern | {temporal.get('duration_type', 'N/A').capitalize()} | Nature of interference events |")
            report_parts.append(f"| Magnitude Level | {temporal.get('magnitude', 'N/A').capitalize()} | Severity of observed anomalies |")
            report_parts.append("")
            subsection += 1

        # Visual Patterns
        if "visual_patterns" in parser_data and parser_data["visual_patterns"]:
            report_parts.append(f"### {section_num}.{subsection} Observed Visual Patterns")
            report_parts.append("")
            report_parts.append("**Summary:** The following patterns were identified through visual inspection of the monitoring data:")
            report_parts.append("")
            patterns = parser_data["visual_patterns"]
            if isinstance(patterns, list):
                for i, pattern in enumerate(patterns, 1):
                    report_parts.append(f"{i}. {pattern}")
            else:
                report_parts.append(f"- {str(patterns)}")
            report_parts.append("")

    def _add_extracted_data_table_v2(self, report_parts: List[str], table_data: dict,
                                     section_num: int) -> None:
        """添加从文本提取的数据表格，按类别分组，带小标题编号"""
        # 按前缀分组数据
        groups = {}
        for key, value in table_data.items():
            if " - " in key:
                prefix, suffix = key.split(" - ", 1)
                if prefix not in groups:
                    groups[prefix] = {}
                groups[prefix][suffix] = value
            else:
                if "General" not in groups:
                    groups["General"] = {}
                groups["General"][key] = value

        # 输出分组表格
        subsection = 1
        for group_name, group_data in groups.items():
            if not group_data:
                continue
            report_parts.append(f"### {section_num}.{subsection} {group_name}")
            report_parts.append("")
            report_parts.append(f"**Summary:** Key parameters extracted for {group_name.lower()} analysis.")
            report_parts.append("")
            report_parts.append("| Parameter | Value |")
            report_parts.append("|-----------|-------|")
            for key, value in group_data.items():
                # 清理值中的字典/列表表示
                if isinstance(value, str) and value.startswith("{"):
                    value = "See detailed analysis"
                report_parts.append(f"| {key} | {value} |")
            report_parts.append("")
            subsection += 1

    def _clean_agent_output(self, output: str, agent_type: str) -> str:
        """
        清理 agent 输出，移除重复内容和格式问题。

        Args:
            output: 原始输出
            agent_type: agent 类型 (analysis/review/report)

        Returns:
            清理后的输出
        """
        if not output:
            return ""

        cleaned = output.strip()

        # 移除所有代码块标记（包括 ```json, ```markdown 等）
        cleaned = re.sub(r'```(?:json|markdown|text)?\\s*', '', cleaned)

        # 移除独立的 JSON 对象（不在代码块中的）
        cleaned = re.sub(r'^\\s*\\{[\\s\\S]*?\\}\\s*$', '', cleaned, flags=re.MULTILINE)
        # 移除多行 JSON 对象
        cleaned = re.sub(r'\\{\\s*"[^"]+\"\\s*:\\s*\\{[\\s\\S]*?\\}\\s*\\}', '', cleaned)
        cleaned = re.sub(r'\\{\\s*"basic_info"[\\s\\S]*?\\}\\s*\\}', '', cleaned)
        # 移除残留的单独大括号
        cleaned = re.sub(r'^\\s*[\\{\\}]\\s*$', '', cleaned, flags=re.MULTILINE)

        # 移除重复的章节标题
        headers_to_remove = [
            r'^#+\\s*ITU\\s*Interference Analysis Report\\s*$',
            r'^#+\\s*Interference Analysis Report\\s*$',
            r'^#+\\s*\\d+\\.\\s*Basic Information\\s*$',
            r'^#+\\s*\\d+\\.\\s*Data Analysis\\s*$',
            r'^#+\\s*\\d+\\.\\s*Compliance Assessment\\s*$',
            r'^#+\\s*\\d+\\.\\s*Conclusions.*$',
            r'^#+\\s*\\d+\\.\\s*Appendix.*$',
            r'^#+\\s*Data Analysis Results\\s*$',
            r'^#+\\s*Compliance Review\\s*$',
            r'^#+\\s*Report Overview\\s*$',
            r'^#+\\s*Data Extraction Results\\s*$',
            r'^#+\\s*Technical Analysis\\s*$',
            r'^#+\\s*ITU Compliance Assessment\\s*$',
        ]
        for pattern in headers_to_remove:
            cleaned = re.sub(pattern, '', cleaned, flags=re.MULTILINE | re.IGNORECASE)

        # 移除只有表头没有内容的表格
        cleaned = re.sub(r'\\|\\s*Field\\s*\\|\\s*Value\\s*\\|\\s*\\n\\|[-\\s|]+\\|\\s*\\n(?!\\|)', '', cleaned)
        cleaned = re.sub(r'\\|\\s*Item\\s*\\|\\s*Value\\s*\\|\\s*\\n\\|[-\\s|]+\\|\\s*\\n(?!\\|)', '', cleaned)
        cleaned = re.sub(r'\\|\\s*Item\\s*\\|\\s*Description\\s*\\|\\s*\\n\\|[-\\s|]+\\|\\s*\\n(?!\\|)', '', cleaned)

        # 移除 "End with:" 等提示文本
        cleaned = re.sub(r'^End with:\\s*$', '', cleaned, flags=re.MULTILINE)

        # 移除占位符文本
        placeholder_patterns = [
            r'\\[1 paragraph explaining.*?\\]',
            r'\\[Visual cue \\d+ with specifics\\]',
            r'\\[.*?placeholder.*?\\]',
            r'<\\|observation\\|>',
            r'<\\|.*?\\|>',
        ]
        for pattern in placeholder_patterns:
            cleaned = re.sub(pattern, '', cleaned, flags=re.IGNORECASE)

        # 处理行
        lines = cleaned.split('\\n')
        result_lines = []
        list_items = []
        in_list = False

        for line in lines:
            stripped = line.strip()

            # 跳过空行（在开头）
            if not result_lines and not stripped:
                continue

            # 检测列表项
            is_list_item = stripped.startswith('- ') or stripped.startswith('* ') or re.match(r'^\\d+\\.\\s', stripped)

            if is_list_item:
                in_list = True
                # 提取列表内容
                if stripped.startswith('- ') or stripped.startswith('* '):
                    content = stripped[2:]
                else:
                    content = re.sub(r'^\\d+\\.\\s*', '', stripped)
                list_items.append(content)
            else:
                # 如果之前在列表中，现在不是列表项了
                if in_list and list_items:
                    # 如果列表项包含 "Key: Value" 格式，转换为表格
                    if all(':' in item for item in list_items[:3]) and len(list_items) >= 2:
                        result_lines.append("")
                        result_lines.append("| Item | Description |")
                        result_lines.append("|------|-------------|")
                        for item in list_items:
                            if ':' in item:
                                key, val = item.split(':', 1)
                                result_lines.append(f"| {key.strip()} | {val.strip()} |")
                            else:
                                result_lines.append(f"| - | {item} |")
                        result_lines.append("")
                    else:
                        # 保持为列表，但使用数字编号
                        result_lines.append("")
                        for i, item in enumerate(list_items, 1):
                            result_lines.append(f"{i}. {item}")
                        result_lines.append("")
                    list_items = []
                    in_list = False

                if stripped:
                    result_lines.append(line)
                elif result_lines and result_lines[-1].strip():
                    result_lines.append("")

        # 处理最后的列表
        if list_items:
            if all(':' in item for item in list_items[:3]) and len(list_items) >= 2:
                result_lines.append("")
                result_lines.append("| Item | Description |")
                result_lines.append("|------|-------------|")
                for item in list_items:
                    if ':' in item:
                        key, val = item.split(':', 1)
                        result_lines.append(f"| {key.strip()} | {val.strip()} |")
                    else:
                        result_lines.append(f"| - | {item} |")
                result_lines.append("")
            else:
                result_lines.append("")
                for i, item in enumerate(list_items, 1):
                    result_lines.append(f"{i}. {item}")
                result_lines.append("")

        cleaned = '\\n'.join(result_lines)

        # 清理多余空行
        while '\\n\\n\\n' in cleaned:
            cleaned = cleaned.replace('\\n\\n\\n', '\\n\\n')

        return cleaned.strip()

    def _aggregate_report(self, parser_output: str, analysis_output: str,
                         review_output: str, report_output: str) -> str:
        """聚合所有 agent 输出为最终报告"""
        report_parts = []

        # 添加报告标题和头部信息
        report_parts.append("# ITU Interference Analysis Report")
        report_parts.append("")
        report_parts.append("---")
        report_parts.append("")
        report_parts.append(f"**Report Date:** {self.current_date}")
        report_parts.append("")
        report_parts.append("**Organization:** Institute of Space Internet, Fudan University")
        report_parts.append("")
        report_parts.append("**Analysis System:** Calference - Satellite Interference Analysis Platform")
        report_parts.append("")
        report_parts.append(f"**Document ID:** ITU-RPT-{self.current_date.replace('-', '')}-{self.image_info.get('constellation', 'SAT').upper()[:3]}")
        report_parts.append("")
        report_parts.append("---")
        report_parts.append("")

        # Section 1: Executive Summary
        report_parts.append("## 1. Executive Summary")
        report_parts.append("")
        report_parts.append(f"**This report presents the interference analysis results for the {self.image_info.get('constellation', 'N/A').upper()} constellation's {self.image_info.get('terminal_type', 'system')}.** The analysis focuses on {self.image_info.get('analysis_type', 'interference metrics')} measurements collected over a 24-hour monitoring period.")
        report_parts.append("")
        report_parts.append("The assessment is conducted in accordance with ITU Radio Regulations and relevant ITU-R Recommendations to ensure compliance with international spectrum management requirements.")
        report_parts.append("")

        section_num = 2

        # Section 2: System Information
        report_parts.append(f"## {section_num}. System Under Test")
        report_parts.append("")
        report_parts.append(f"### {section_num}.1 Basic Information")
        report_parts.append("")
        report_parts.append("| Parameter | Value |")
        report_parts.append("|-----------|-------|")
        report_parts.append(f"| Constellation Name | {self.image_info.get('constellation', 'N/A').upper()} |")
        report_parts.append(f"| Terminal Type | {self.image_info.get('terminal_type', 'N/A')} |")
        report_parts.append(f"| Analysis Metric | {self.image_info.get('analysis_type', 'N/A')} |")
        report_parts.append(f"| Monitoring Period | 24 hours (00:00 - 23:59 UTC) |")
        report_parts.append(f"| Source Data File | {self.image_info.get('filename', 'N/A')} |")
        report_parts.append("")

        report_parts.append(f"### {section_num}.2 Applicable ITU Standards")
        report_parts.append("")
        report_parts.append("The following ITU-R Recommendations are referenced in this analysis:")
        report_parts.append("")
        report_parts.append("| Standard | Description |")
        report_parts.append("|----------|-------------|")
        report_parts.append("| ITU-R S.1503 | Functional description to be used in developing software tools for determining conformity of non-GSO FSS systems |")
        report_parts.append("| ITU-R S.1325 | Simulation methodologies for determining statistics of short-term interference |")
        report_parts.append("| ITU-R S.1528 | Satellite antenna radiation patterns for non-GSO orbit systems |")
        report_parts.append("| ITU-R SF.1395 | Minimum propagation attenuation due to atmospheric gases for frequency coordination |")
        report_parts.append("| ITU-R P.618 | Propagation data and prediction methods for earth-space telecommunication systems |")
        report_parts.append("")

        section_num += 1

        # Section 3: Data Analysis (from parser output)
        if parser_output:
            report_parts.append(f"## {section_num}. Measurement Data Analysis")
            report_parts.append("")
            report_parts.append(f"**This section presents the quantitative analysis of the {self.image_info.get('analysis_type', 'interference')} measurements extracted from the monitoring data.**")
            report_parts.append("")

            subsection = 1
            try:
                # 尝试解析 JSON
                parser_data = json.loads(parser_output)
                self._add_structured_data_section_v2(report_parts, parser_data, section_num, subsection)
            except (json.JSONDecodeError, KeyError, TypeError) as e:
                print(f"[DEBUG] JSON parsing failed: {e}, converting raw output to table format")
                # 尝试从原始文本中提取键值对
                table_data = self._parse_text_to_table_data(parser_output)
                if table_data:
                    self._add_extracted_data_table_v2(report_parts, table_data, section_num)
                else:
                    report_parts.append(parser_output)
                    report_parts.append("")

            section_num += 1

        # Section 4: Technical Analysis (from analysis agent)
        if analysis_output:
            analysis_clean = self._clean_agent_output(analysis_output, "analysis")
            if analysis_clean.strip():
                report_parts.append(f"## {section_num}. Technical Analysis")
                report_parts.append("")
                report_parts.append(f"**The following technical analysis evaluates the interference characteristics and potential impact on system performance.**")
                report_parts.append("")
                report_parts.append(analysis_clean)
                report_parts.append("")
                section_num += 1

        # Section 5: Compliance Assessment (from review agent)
        if review_output:
            review_clean = self._clean_agent_output(review_output, "review")
            if review_clean.strip():
                report_parts.append(f"## {section_num}. ITU Compliance Assessment")
                report_parts.append("")
                report_parts.append(f"**This section assesses the system's compliance with applicable ITU Radio Regulations and Recommendations.**")
                report_parts.append("")
                report_parts.append(review_clean)
                report_parts.append("")
                section_num += 1

        # Section 6: Conclusions and Recommendations (from report agent)
        if report_output:
            report_clean = self._clean_agent_output(report_output, "report")
            if report_clean.strip():
                report_parts.append(f"## {section_num}. Conclusions and Recommendations")
                report_parts.append("")
                report_parts.append(f"**Based on the comprehensive analysis presented above, the following conclusions and recommendations are provided.**")
                report_parts.append("")
                report_parts.append(report_clean)
                report_parts.append("")
                section_num += 1

        # Section: References
        report_parts.append("---")
        report_parts.append(f"## {section_num}. References")
        report_parts.append("")
        report_parts.append("### ITU-R Recommendations")
        report_parts.append("")
        report_parts.append("1. **ITU-R S.1503-3** (2013): Functional description to be used in developing software tools for determining conformity of non-geostationary-satellite orbit fixed-satellite service systems or networks with limits contained in Article 22 of the Radio Regulations")
        report_parts.append("")
        report_parts.append("2. **ITU-R S.1325-3** (2003): Simulation methodologies for determining statistics of short-term interference between co-frequency, codirectional non-geostationary-satellite orbit fixed-satellite service systems in circular orbits and other non-geostationary fixed-satellite service systems in circular orbits or geostationary-satellite orbit fixed-satellite service networks")
        report_parts.append("")
        report_parts.append("3. **ITU-R S.1528** (2001): Satellite antenna radiation patterns for non-geostationary orbit satellite antennas operating in the fixed-satellite service below 30 GHz")
        report_parts.append("")
        report_parts.append("4. **ITU-R SF.1395** (1999): Minimum propagation attenuation due to atmospheric gases for use in frequency coordination between the fixed-satellite service and the fixed service")
        report_parts.append("")
        report_parts.append("5. **ITU-R P.618-13** (2017): Propagation data and prediction methods required for the design of Earth-space telecommunication systems")
        report_parts.append("")
        section_num += 1

        # Appendix
        report_parts.append("---")
        report_parts.append(f"## {section_num}. Appendix: Report Metadata")
        report_parts.append("")
        report_parts.append("| Parameter | Value |")
        report_parts.append("|-----------|-------|")
        report_parts.append("| Report Type | Multi-Agent Dialogue-Based Analysis |")
        report_parts.append("| Analysis Pipeline | Parser → Analysis → Review → Report |")
        report_parts.append(f"| Generation Date | {self.current_date} |")
        report_parts.append("| Analysis System | Calference v1.0 |")
        report_parts.append("| Organization | Institute of Space Internet, Fudan University |")
        report_parts.append("")
        report_parts.append("---")
        report_parts.append("")
        report_parts.append("*This report was automatically generated by the Calference Satellite Interference Analysis Platform.*")
        report_parts.append("")
        report_parts.append("*© 2026 Institute of Space Internet, Fudan University. All rights reserved.*")
        report_parts.append("")

        # 合并所有部分
        final_report = "\n".join(report_parts)
        # 清理多余空行
        while "\n\n\n" in final_report:
            final_report = final_report.replace("\n\n\n", "\n\n")

        return final_report

    def _create_docx_with_image(self, md_path: str, docx_path: str, image_path: str) -> None:
        """将 Markdown 转换为 DOCX 并插入图片"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            # 尝试使用 pandoc
            pandoc_exe = shutil.which("pandoc")
            if pandoc_exe:
                print("[INFO] Using pandoc to convert markdown to DOCX")
                try:
                    subprocess.run(
                        [pandoc_exe, "--from", "gfm", "--to", "docx", md_path, "-o", docx_path],
                        check=True,
                        capture_output=True,
                        text=True,
                    )
                    doc = Document(docx_path)
                    print("[INFO] Pandoc conversion successful")
                except subprocess.CalledProcessError as e:
                    print(f"[WARNING] Pandoc conversion failed: {e.stderr}")
                    doc = self._create_docx_from_markdown(md_path)
            else:
                print("[INFO] pandoc not found, using python-docx")
                doc = self._create_docx_from_markdown(md_path)

            # 插入图片
            if os.path.exists(image_path):
                print(f"[INFO] Inserting image from: {image_path}")
                img_para = doc.add_paragraph()
                img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = img_para.add_run()
                run.add_picture(image_path, width=Inches(5.5))

                caption_para = doc.add_paragraph()
                caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption_run = caption_para.add_run(
                    f"Figure 1: {self.image_info.get('filename', 'Interference Map')}"
                )
                caption_run.font.size = Pt(10)
                caption_run.font.italic = True

            doc.save(docx_path)
            print(f"[INFO] DOCX saved: {docx_path}")

        except Exception as e:
            print(f"[ERROR] Failed to create DOCX: {e}")
            import traceback
            traceback.print_exc()

    def _create_docx_from_markdown(self, md_path: str) -> "Document":
        """
        使用纯 python-docx 从 markdown 创建 Word 文档。

        Args:
            md_path: Markdown 文件路径

        Returns:
            Document 对象
        """
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement

        doc = Document()

        # 读取 markdown 内容
        with open(md_path, "r", encoding="utf-8") as f:
            content = f.read()

        lines = content.split("\n")
        i = 0
        in_code_block = False
        code_content = []
        in_table = False
        table_rows = []

        print(f"[DEBUG] Processing {len(lines)} lines from markdown")

        while i < len(lines):
            line = lines[i]

            # 处理代码块
            if line.strip().startswith("```"):
                if in_code_block:
                    # 结束代码块
                    in_code_block = False
                    if code_content:
                        # 添加代码块内容
                        code_para = doc.add_paragraph()
                        code_para.style = "Normal"
                        code_run = code_para.add_run("\n".join(code_content))
                        code_run.font.name = "Courier New"
                        code_run.font.size = Pt(9)
                        code_content = []
                else:
                    # 开始代码块
                    in_code_block = True
                i += 1
                continue

            if in_code_block:
                code_content.append(line)
                i += 1
                continue

            # 处理表格
            if line.strip().startswith("|") and "|" in line.strip()[1:]:
                if not in_table:
                    in_table = True
                    table_rows = []
                    print(f"[DEBUG] Table started at line {i}")
                # 跳过分隔行（包含 --- 或 :-- 或 --: 等）
                stripped_line = line.strip()
                # 更精确的分隔行检测：必须包含连续的 - 符号
                if re.match(r"^\|[\s\-:|]+\|$", stripped_line) and "---" in stripped_line:
                    print(f"[DEBUG] Skipping separator line at {i}: {stripped_line[:50]}")
                    i += 1
                    continue
                # 解析表格行
                cells = [cell.strip() for cell in line.strip().split("|")[1:-1]]
                # 过滤掉空行和只有空格的行
                if cells and any(cell.strip() for cell in cells):
                    print(f"[DEBUG] Adding table row at line {i} with {len(cells)} cells")
                    table_rows.append(cells)
                else:
                    print(f"[DEBUG] Skipping empty table row at line {i}")
                i += 1
                continue
            elif in_table:
                # 表格结束，创建表格
                in_table = False
                print(f"[DEBUG] Table ended at line {i}, total rows: {len(table_rows)}")
                if table_rows and len(table_rows) >= 1:  # 至少要有表头
                    self._add_table_to_doc(doc, table_rows)
                    table_rows = []
                else:
                    print(f"[DEBUG] Skipping table creation - insufficient rows")
                # 注意：不要 continue，让当前行继续被处理

            # 处理标题
            if line.startswith("# "):
                heading = doc.add_heading(line[2:].strip(), level=1)
                i += 1
                continue
            elif line.startswith("## "):
                heading = doc.add_heading(line[3:].strip(), level=2)
                i += 1
                continue
            elif line.startswith("### "):
                heading = doc.add_heading(line[4:].strip(), level=3)
                i += 1
                continue
            elif line.startswith("#### "):
                heading = doc.add_heading(line[5:].strip(), level=4)
                i += 1
                continue

            # 处理水平线
            if line.strip() == "---":
                # 添加水平线
                para = doc.add_paragraph()
                para.add_run("─" * 50)
                i += 1
                continue

            # 处理列表项
            if line.strip().startswith("- ") or line.strip().startswith("* "):
                text = line.strip()[2:]
                para = doc.add_paragraph(style="List Bullet")
                self._add_formatted_text(para, text)
                i += 1
                continue

            # 处理数字列表
            match = re.match(r"^\d+\.\s+(.+)$", line.strip())
            if match:
                para = doc.add_paragraph(style="List Number")
                self._add_formatted_text(para, match.group(1))
                i += 1
                continue

            # 处理普通段落
            if line.strip():
                para = doc.add_paragraph()
                self._add_formatted_text(para, line.strip())

            i += 1

        # 处理最后的表格
        if in_table and table_rows and len(table_rows) >= 1:
            print(f"[DEBUG] Processing final table with {len(table_rows)} rows")
            self._add_table_to_doc(doc, table_rows)

        print(f"[DEBUG] Markdown processing complete, total paragraphs: {len(doc.paragraphs)}")
        print(f"[DEBUG] Total tables created: {len(doc.tables)}")

        return doc

    def _add_table_to_doc(self, doc: "Document", rows: List[List[str]]) -> None:
        """
        向文档添加专业格式的表格。

        Args:
            doc: Document 对象
            rows: 表格行数据
        """
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import nsdecls
        from docx.oxml import parse_xml

        if not rows or len(rows) < 1:
            print("[DEBUG] No rows to add to table")
            return

        num_cols = len(rows[0]) if rows else 0

        if num_cols == 0:
            print("[DEBUG] Column count is 0")
            return

        print(f"[DEBUG] Creating table with {len(rows)} rows and {num_cols} columns")

        # 标准化所有行的列数
        normalized_rows = []
        for row in rows:
            if len(row) < num_cols:
                # 补齐缺失的列
                row = row + [""] * (num_cols - len(row))
            elif len(row) > num_cols:
                # 截断多余的列
                row = row[:num_cols]
            normalized_rows.append(row)

        try:
            # 创建表格
            table = doc.add_table(rows=len(normalized_rows), cols=num_cols)
            table.style = "Table Grid"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # 设置表格属性
            tbl = table._element
            tblPr = tbl.tblPr
            if tblPr is None:
                tblPr = parse_xml(r'<w:tblPr {}/>', nsdecls('w'))
                tbl.insert(0, tblPr)

            # 设置表格宽度为页面宽度
            tblW = parse_xml(
                r'<w:tblW {} w:w="5000" w:type="pct"/>'.format(nsdecls('w'))
            )
            tblPr.append(tblW)

            # 设置表格边框（所有边都有黑色边框）
            tblBorders = parse_xml(
                r'<w:tblBorders {}>'
                r'<w:top w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
                r'<w:left w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
                r'<w:bottom w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
                r'<w:right w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
                r'<w:insideH w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
                r'<w:insideV w:val="single" w:sz="12" w:space="0" w:color="000000"/>'
                r'</w:tblBorders>'.format(nsdecls('w'))
            )
            tblPr.append(tblBorders)

            # 设置表格单元格间距
            tblCellMar = parse_xml(
                r'<w:tblCellMar {}>'
                r'<w:top w:w="50" w:type="dxa"/>'
                r'<w:left w:w="50" w:type="dxa"/>'
                r'<w:bottom w:w="50" w:type="dxa"/>'
                r'<w:right w:w="50" w:type="dxa"/>'
                r'</w:tblCellMar>'.format(nsdecls('w'))
            )
            tblPr.append(tblCellMar)

            # 设置表格阴影效果
            tblShd = parse_xml(
                r'<w:tblShd {} w:val="clear" w:color="auto" w:fill="FFFFFF"/>'.format(nsdecls('w'))
            )
            tblPr.append(tblShd)

            # 添加表格布局设置
            tblLayout = parse_xml(
                r'<w:tblLayout {} w:type="auto"/>'.format(nsdecls('w'))
            )
            tblPr.append(tblLayout)

            # 计算每列的宽度（均匀分配）
            col_width = Inches(5.5 / num_cols) if num_cols > 0 else Inches(1)

            # 设置列宽
            for row in table.rows:
                for idx, cell in enumerate(row.cells):
                    if idx < num_cols:
                        tcPr = cell._element.get_or_add_tcPr()
                        col_width_value = int(5000 / num_cols)
                        tcW = parse_xml(
                            r'<w:tcW {} w:w="{}" w:type="dxa"/>'.format(nsdecls('w'), col_width_value)
                        )
                        # 移除旧的宽度设置
                        for tcW_old in tcPr.findall('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}tcW'):
                            tcPr.remove(tcW_old)
                        tcPr.insert(0, tcW)

            # 填充表格
            for i, row_data in enumerate(normalized_rows):
                row = table.rows[i]

                # 设置行高
                row.height = Pt(32)

                for j, cell_text in enumerate(row_data):
                    if j < num_cols and j < len(row.cells):
                        cell = row.cells[j]

                        # 设置列宽
                        tcPr = cell._element.get_or_add_tcPr()
                        tcW = parse_xml(
                            r'<w:tcW {} w:w="1000" w:type="auto"/>'.format(nsdecls('w'))
                        )
                        tcPr.append(tcW)

                        # 设置单元格内边距
                        tcMar = parse_xml(
                            r'<w:tcMar {}>'
                            r'<w:top w:w="100" w:type="dxa"/>'
                            r'<w:left w:w="100" w:type="dxa"/>'
                            r'<w:bottom w:w="100" w:type="dxa"/>'
                            r'<w:right w:w="100" w:type="dxa"/>'
                            r'</w:tcMar>'.format(nsdecls('w'))
                        )
                        tcPr.append(tcMar)

                        # 设置单元格垂直对齐方式为居中
                        vAlign = parse_xml(
                            r'<w:vAlign {} w:val="center"/>'.format(nsdecls('w'))
                        )
                        tcPr.append(vAlign)

                        # 清空默认段落
                        cell.text = ""
                        para = cell.paragraphs[0]

                        # 设置段落对齐方式
                        if i == 0:
                            # 表头居中对齐
                            para.alignment = 1  # WD_ALIGN_PARAGRAPH.CENTER
                        else:
                            # 数据行左对齐
                            para.alignment = 0  # WD_ALIGN_PARAGRAPH.LEFT

                        # 设置段落行间距
                        para.paragraph_format.line_spacing = 1.15
                        para.paragraph_format.space_before = Pt(0)
                        para.paragraph_format.space_after = Pt(0)

                        # 处理 markdown 加粗语法 **text**
                        self._add_formatted_cell_text(para, cell_text if cell_text else "", Pt(11), is_header=(i == 0))

                # 设置表头行背景色和文字颜色
                if i == 0:
                    for j in range(min(num_cols, len(row.cells))):
                        cell = row.cells[j]
                        try:
                            # 设置深蓝色背景 (4472C4)
                            tcPr = cell._element.get_or_add_tcPr()
                            shading_elm = parse_xml(
                                r'<w:shd {} w:fill="4472C4"/>'.format(nsdecls('w'))
                            )
                            tcPr.append(shading_elm)

                            # 设置表头单元格边框为深色
                            tcBorders = parse_xml(
                                r'<w:tcBorders {}>'
                                r'<w:top w:val="single" w:sz="12" w:space="0" w:color="2F5496"/>'
                                r'<w:left w:val="single" w:sz="12" w:space="0" w:color="2F5496"/>'
                                r'<w:bottom w:val="single" w:sz="12" w:space="0" w:color="2F5496"/>'
                                r'<w:right w:val="single" w:sz="12" w:space="0" w:color="2F5496"/>'
                                r'</w:tcBorders>'.format(nsdecls('w'))
                            )
                            tcPr.append(tcBorders)

                            # 设置表头文字为白色、加粗、增大字体
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.color.rgb = RGBColor(255, 255, 255)
                                    run.font.bold = True
                                    run.font.size = Pt(12)
                        except Exception as e:
                            print(f"[WARNING] Failed to set header cell style: {e}")
                else:
                    # 设置数据行的交替背景色（斑马纹）和边框
                    for j in range(min(num_cols, len(row.cells))):
                        cell = row.cells[j]
                        try:
                            tcPr = cell._element.get_or_add_tcPr()

                            # 设置交替背景色
                            if i % 2 == 1:
                                # 设置浅灰色背景 (F2F2F2)
                                shading_elm = parse_xml(
                                    r'<w:shd {} w:fill="F2F2F2"/>'.format(nsdecls('w'))
                                )
                            else:
                                # 白色背景
                                shading_elm = parse_xml(
                                    r'<w:shd {} w:fill="FFFFFF"/>'.format(nsdecls('w'))
                                )
                            tcPr.append(shading_elm)

                            # 设置数据行单元格边框为浅灰色
                            tcBorders = parse_xml(
                                r'<w:tcBorders {}>'
                                r'<w:top w:val="single" w:sz="8" w:space="0" w:color="D0D0D0"/>'
                                r'<w:left w:val="single" w:sz="8" w:space="0" w:color="D0D0D0"/>'
                                r'<w:bottom w:val="single" w:sz="8" w:space="0" w:color="D0D0D0"/>'
                                r'<w:right w:val="single" w:sz="8" w:space="0" w:color="D0D0D0"/>'
                                r'</w:tcBorders>'.format(nsdecls('w'))
                            )
                            tcPr.append(tcBorders)

                            # 设置数据行文字颜色为深灰色
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    run.font.color.rgb = RGBColor(64, 64, 64)
                        except Exception as e:
                            print(f"[WARNING] Failed to set data row style: {e}")

            # 添加空行
            spacer = doc.add_paragraph()
            spacer.paragraph_format.space_before = Pt(6)
            spacer.paragraph_format.space_after = Pt(6)

            print(f"[DEBUG] Table added successfully")

        except Exception as e:
            print(f"[ERROR] Failed to create table: {e}")
            import traceback
            traceback.print_exc()

    def _add_formatted_cell_text(self, para, text: str, font_size, is_header: bool = False) -> None:
        """
        向单元格段落添加格式化文本，处理 markdown 加粗语法。

        Args:
            para: 段落对象
            text: 文本内容
            font_size: 字体大小
            is_header: 是否为表头行
        """
        from docx.shared import RGBColor

        # 处理空值
        if text is None:
            text = ""
        text = str(text).strip()

        # 如果文本为空，添加一个空格以保持单元格结构
        if not text:
            run = para.add_run(" ")
            run.font.size = font_size
            run.font.name = "Calibri"
            if is_header:
                run.font.color.rgb = RGBColor(255, 255, 255)
            return

        # 处理 markdown 加粗语法 **text**
        pattern = r'(\*\*(.+?)\*\*|([^*]+))'
        matches = re.findall(pattern, text)

        # 如果没有匹配到任何内容，直接添加原始文本
        if not matches:
            run = para.add_run(text)
            run.font.size = font_size
            run.font.name = "Calibri"
            if is_header:
                run.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
            else:
                run.font.color.rgb = RGBColor(64, 64, 64)
            return

        for match in matches:
            full, bold_text, normal_text = match
            if bold_text:
                run = para.add_run(bold_text)
                run.bold = True
                run.font.size = font_size
                run.font.name = "Calibri"
                if is_header:
                    run.font.color.rgb = RGBColor(255, 255, 255)
                else:
                    run.font.color.rgb = RGBColor(0, 0, 0)
            elif normal_text:
                run = para.add_run(normal_text)
                run.font.size = font_size
                run.font.name = "Calibri"
                if is_header:
                    run.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                else:
                    run.font.color.rgb = RGBColor(64, 64, 64)

    def _add_formatted_text(self, para, text: str) -> None:
        """
        向段落添加格式化文本，处理粗体和斜体。

        Args:
            para: 段落对象
            text: 文本内容
        """
        # 简单处理粗体 **text** 和斜体 *text*
        pattern = r"(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`|([^*`]+))"
        matches = re.findall(pattern, text)

        for match in matches:
            full, bold, italic, code, normal = match
            if bold:
                run = para.add_run(bold)
                run.bold = True
            elif italic:
                run = para.add_run(italic)
                run.italic = True
            elif code:
                run = para.add_run(code)
                run.font.name = "Courier New"
            elif normal:
                para.add_run(normal)


# ============================================================================
# 主函数
# ============================================================================
async def main(image_path: str) -> None:
    """主入口"""
    configure_proxies()

    print("==== ITU Interference Report Generation (Standalone) ====")
    print(f"[INFO] Image path: {image_path}")
    print(f"[INFO] Model: {LLM_MODEL_NAME}")
    print(f"[INFO] API URL: {LLM_BASE_URL}")

    if not os.path.exists(image_path):
        print(f"[ERROR] Image file not found: {image_path}")
        sys.exit(1)

    # 输出目录为图片所在目录
    output_dir = os.path.dirname(os.path.abspath(image_path))

    filename = os.path.basename(image_path)
    image_info = parse_image_info(filename)
    current_date = datetime.now().strftime("%Y-%m-%d")

    pipeline = StandaloneAnalysisPipeline(image_info, current_date, output_dir)
    await pipeline.run(image_path)


if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python standalone_itu_analyzer.py <image_path>")
        print("Example: python standalone_itu_analyzer.py /path/to/interference_map.png")
        sys.exit(1)

    image_path = sys.argv[1]
    asyncio.run(main(image_path))
