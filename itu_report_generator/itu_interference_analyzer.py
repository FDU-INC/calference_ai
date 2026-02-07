# GNU GENERAL PUBLIC LICENSE
# Version 3, 29 June 2007
#
# Copyright (C) 2025 FDU-INC
#
# This program is free software: you can redistribute it and/or modify
# it under the terms of the GNU General Public License as published by
# the Free Software Foundation, either version 3 of the License, or
# (at your option) any later version.
#
# This program is distributed in the hope that it will be useful,
# but WITHOUT ANY WARRANTY; without even the implied warranty of
# MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
# GNU General Public License for more details.
#
# You should have received a copy of the GNU General Public License
# along with this program.  If not, see <https://www.gnu.org/licenses/>.
#
# Author: yjh
# Date: 2026-02-07
# Description: ITU Interference Analyzer - Multi-Agent Report Generation System
#              Automated interference analysis for satellite communication systems
#              using a multi-agent dialogue architecture with ITU compliance assessment.

import os
import asyncio
import json
from datetime import datetime
from typing import Dict, List
import shutil
import subprocess

from PIL import Image
from autogen_core import Image as AutoGenImage
from autogen_agentchat.agents import AssistantAgent
from autogen_agentchat.conditions import TextMentionTermination
from autogen_agentchat.teams import RoundRobinGroupChat
from autogen_agentchat.messages import MultiModalMessage
from autogen_ext.models.openai import OpenAIChatCompletionClient

from config import (
    LLM_MODEL_NAME,
    LLM_BASE_URL,
    LLM_API_KEY,
    INPUT_IMAGE_DIR,
    OUTPUT_REPORT_DIR,
    OUTPUT_REPORTS_DIR,
    OUTPUT_METADATA_DIR,
    OUTPUT_DATAFLOW_DIR,
)
from agent_config import (
    AgentConfigBuilder,
    DataFlowManager,
    AgentSequenceValidator,
)


def configure_proxies() -> None:
    """Configure HTTP proxy"""
    http_proxy = "http://10.192.54.148:7897"
    os.environ["HTTP_PROXY"] = http_proxy
    os.environ["HTTPS_PROXY"] = http_proxy
    print(f"[INFO] HTTP(S) proxy configured: {http_proxy}")


def compress_image(image_path: str, max_size: int = 800, quality: int = 80) -> str:
    """Compress image to reduce token consumption"""
    img = Image.open(image_path)
    width, height = img.size
    if max(width, height) > max_size:
        scale = max_size / max(width, height)
        new_width = int(width * scale)
        new_height = int(height * scale)
        img = img.resize((new_width, new_height), Image.Resampling.LANCZOS)
        print(f"[INFO] Image resized: {width}x{height} -> {new_width}x{new_height}")

    compressed_path = image_path.replace(".png", "_compressed.jpg")
    img.convert("RGB").save(compressed_path, "JPEG", quality=quality, optimize=True)

    original_size = os.path.getsize(image_path) / 1024
    compressed_size = os.path.getsize(compressed_path) / 1024
    print(
        f"[INFO] Image compressed: {original_size:.1f}KB -> {compressed_size:.1f}KB (saved {100*(1-compressed_size/original_size):.1f}%)"
    )

    return compressed_path


def parse_image_info(filename: str) -> Dict:
    """Parse image information from filename"""
    parts = filename.split("_")
    constellation = parts[0] if parts else "Unknown"
    terminal_type = "Ground Terminal System" if "earth" in filename.lower() else "Satellite System"

    analysis_types_map = {
        "cinr": "CINR",
        "cir": "CIR",
        "cnr": "CNR",
        "epfd": "EPFD",
        "inr": "INR",
        "link_count": "Link Count",
        "pfd": "PFD",
        "temp": "Delta T/T",
    }
    analysis_type = next(
        (v for k, v in analysis_types_map.items() if k in filename.lower()), "Unknown"
    )

    return {
        "constellation": constellation,
        "terminal_type": terminal_type,
        "analysis_type": analysis_type,
        "filename": filename,
    }


def create_model_client(
    max_tokens: int = 2048, temperature: float = 0.3
) -> OpenAIChatCompletionClient:
    """Create LLM client"""
    base_url = LLM_BASE_URL.rstrip("/")
    if base_url.endswith("/v1"):
        base_url = base_url[:-3]

    return OpenAIChatCompletionClient(
        model=LLM_MODEL_NAME,
        base_url=base_url,
        api_key=LLM_API_KEY,
        model_capabilities={
            "vision": True,
            "function_calling": True,
            "json_output": True,
        },
        create_args={
            "max_tokens": max_tokens,
            "temperature": temperature,
        },
    )


class ConfigurableDialoguePipeline:
    """
    Configurable dialogue-based multi-agent pipeline

    Features:
    1. Fully configuration-based
    2. Support dynamic agent sequences
    3. Complete data flow tracking
    4. Detailed audit logging
    5. Image integration in reports
    """

    def __init__(self, image_info: dict, current_date: str, image_path: str = None):
        self.image_info = image_info
        self.current_date = current_date
        self.image_path = image_path
        self.data_flow_manager = DataFlowManager()
        self.agents = {}
        self.agent_configs = {}

    def _create_agents_from_configs(self) -> Dict[str, AssistantAgent]:
        """Create agents from configurations"""
        configs = AgentConfigBuilder.build_all_configs()
        agents = {}

        for name, config in configs.items():
            client = create_model_client(
                max_tokens=config.max_tokens,
                temperature=config.temperature,
            )
            agent = AssistantAgent(
                config.name,
                model_client=client,
                description=config.description,
                system_message=config.system_message,
            )
            agents[name] = agent
            self.agent_configs[name] = config

        return agents

    async def run(
        self,
        image_path: str,
        compress_img: bool = False,
    ) -> str:
        """
        运行管道

        处理流程：
        1. 加载图像
        2. 从配置创建 Agent
        3. 验证 Agent 序列
        4. 启动对话
        5. 追踪数据流
        6. 保存结果和审计日志
        """

        # 1. Load image
        if compress_img:
            print("[INFO] Compressing image...")
            working_image_path = compress_image(image_path, max_size=800, quality=80)
        else:
            print("[INFO] Using original image")
            working_image_path = image_path

        print(f"[INFO] Loading image: {working_image_path}")
        pil_image = Image.open(working_image_path)
        autogen_image = AutoGenImage(pil_image)

        # 2. Create agents
        self.agents = self._create_agents_from_configs()

        # 3. Validate agent sequence
        configs_list = list(self.agent_configs.values())
        if not AgentSequenceValidator.validate_sequence(configs_list):
            print("[ERROR] Agent sequence validation failed")
            return ""

        AgentSequenceValidator.print_sequence_info(configs_list)

        # 4. Record initial input
        self.data_flow_manager.record_input(
            "parser_agent",
            f"Image: {os.path.basename(image_path)}",
            "image",
        )

        # 5. Build initial prompt
        initial_prompt = f"""Context Information:
- Constellation: {self.image_info['constellation']}
- Terminal Type: {self.image_info['terminal_type']}
- Analysis Type: {self.image_info['analysis_type']}
- Filename: {self.image_info['filename']}
- Date: {self.current_date}

IMPORTANT: This is a heat map / time-series plot showing interference metric values.
Look carefully at the COLORBAR (legend) on the side of the figure to identify min/max values.

Please analyze this interference map and extract structured data."""

        # 6. Create GroupChat
        print("[INFO] Starting configurable dialogue-based multi-agent pipeline...")
        termination = TextMentionTermination(text="[REPORT_DONE]")

        # Arrange agents in config order
        agent_sequence = [
            self.agents["parser"],
            self.agents["analysis"],
            self.agents["review"],
            self.agents["report"],
        ]

        group_chat = RoundRobinGroupChat(
            agent_sequence,
            termination_condition=termination,
        )

        # 7. Start dialogue
        task = MultiModalMessage(content=[initial_prompt, autogen_image], source="user")

        try:
            print("[INFO] Parser Agent analyzing image...")
            result = await group_chat.run(task=task)
            messages = result.messages

            # 8. Debug: Print message structure
            print(f"[DEBUG] Total messages: {len(messages)}")
            for i, msg in enumerate(messages):
                source = getattr(msg, "source", "N/A")
                agent_name = getattr(msg, "agent_name", "N/A")
                name = getattr(msg, "name", "N/A")
                print(f"[DEBUG] Message {i}: type={type(msg).__name__}, source={source}, agent_name={agent_name}, name={name}")
                if hasattr(msg, "content") and isinstance(msg.content, str):
                    has_terminator = any(
                        marker in msg.content
                        for marker in ["[PARSER_DONE]", "[ANALYSIS_DONE]", "[REVIEW_DONE]", "[REPORT_DONE]"]
                    )
                    print(f"[DEBUG]   content_length={len(msg.content)}, has_terminator={has_terminator}")

            # 9. Track data flow
            self._track_data_flow(messages)

            # 10. Extract final report
            final_report = self._extract_final_report(messages)

            if not final_report:
                print("[WARNING] No report generated")
                return ""

            print(f"[INFO] Report generated: {len(final_report)} chars")

            # 11. Save results to organized directories
            os.makedirs(OUTPUT_REPORTS_DIR, exist_ok=True)
            os.makedirs(OUTPUT_METADATA_DIR, exist_ok=True)
            os.makedirs(OUTPUT_DATAFLOW_DIR, exist_ok=True)

            # Save markdown report to reports/
            md_path = os.path.join(OUTPUT_REPORTS_DIR, "interference_report_en.md")
            with open(md_path, "w", encoding="utf-8") as f:
                f.write(final_report)
            print(f"[INFO] Markdown saved: {md_path}")

            # Save data flow record to data_flow/
            flow_path = os.path.join(OUTPUT_DATAFLOW_DIR, "data_flow.json")
            self.data_flow_manager.export_data_flow(flow_path)
            print(f"[INFO] Data flow saved: {flow_path}")

            # Save audit log to metadata/
            audit_path = os.path.join(OUTPUT_METADATA_DIR, "audit_log.json")
            self.data_flow_manager.export_audit_log(audit_path)
            print(f"[INFO] Audit log saved: {audit_path}")

            # Save agent configs to metadata/
            config_path = os.path.join(OUTPUT_METADATA_DIR, "agent_configs.json")
            AgentConfigBuilder.export_configs(config_path)
            print(f"[INFO] Agent configs saved: {config_path}")

            # Save data flow summary to data_flow/
            summary_path = os.path.join(OUTPUT_DATAFLOW_DIR, "data_flow_summary.json")
            with open(summary_path, "w", encoding="utf-8") as f:
                json.dump(
                    self.data_flow_manager.get_data_flow_summary(),
                    f,
                    indent=2,
                    ensure_ascii=False,
                )
            print(f"[INFO] Data flow summary saved: {summary_path}")

            # 12. Convert to Word with image integration and save to reports/
            docx_path = os.path.join(OUTPUT_REPORTS_DIR, "interference_report_en.docx")
            self._create_docx_with_image(md_path, docx_path, image_path)
            print(f"[INFO] DOCX saved: {docx_path}")

            # 13. Cleanup
            if compress_img and working_image_path != image_path and os.path.exists(working_image_path):
                os.remove(working_image_path)
                print(f"[INFO] Cleaned up compressed image")

            print("==== Report generation finished ====")
            return final_report

        except Exception as e:
            print(f"[ERROR] Pipeline failed: {e}")
            import traceback

            traceback.print_exc()
            return ""

    def _create_docx_with_image(self, md_path: str, docx_path: str, image_path: str) -> None:
        """Convert markdown to DOCX and integrate image"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            # First try to convert markdown to docx using pandoc
            pandoc_exe = shutil.which("pandoc")
            if pandoc_exe:
                subprocess.run(
                    [pandoc_exe, "--from", "gfm", "--to", "docx", md_path, "-o", docx_path],
                    check=True,
                )
                doc = Document(docx_path)
            else:
                # 没有 pandoc，使用纯 python-docx 创建文档
                print("[INFO] pandoc not found, using python-docx to create DOCX")
                doc = self._create_docx_from_markdown(md_path)

            # 插入图片
            if os.path.exists(image_path):
                # Find the position to insert image (after "Data Analysis" section header)
                insert_index = -1
                for i, para in enumerate(doc.paragraphs):
                    # 查找 "Data Analysis" 或 "Basic Information" 标题
                    if "Data Analysis" in para.text or "Basic Information" in para.text:
                        insert_index = i + 1
                        break

                # 如果没找到，在报告标题后插入
                if insert_index == -1:
                    for i, para in enumerate(doc.paragraphs):
                        if "Interference Analysis Report" in para.text:
                            insert_index = i + 1
                            break

                # 如果还是没找到，在开头插入
                if insert_index == -1:
                    insert_index = 1

                # 创建图片段落
                img_para = doc.add_paragraph()
                img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = img_para.add_run()
                run.add_picture(image_path, width=Inches(5.5))

                # 创建图片标题段落
                caption_para = doc.add_paragraph()
                caption_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption_run = caption_para.add_run(
                    f"Figure 1: {self.image_info.get('filename', 'Interference Map')}"
                )
                caption_run.font.size = Pt(10)
                caption_run.font.italic = True

                # 添加空行
                spacer_para = doc.add_paragraph()

                # 将新创建的段落移动到正确位置
                if insert_index < len(doc.paragraphs) - 3:
                    target_element = doc.paragraphs[insert_index]._element
                    # 移动三个新段落到目标位置之前
                    target_element.addprevious(img_para._element)
                    target_element.addprevious(caption_para._element)
                    target_element.addprevious(spacer_para._element)

                print(f"[INFO] Image integrated into DOCX")
            else:
                print(f"[WARNING] Image file not found: {image_path}")

            doc.save(docx_path)
            print(f"[INFO] DOCX saved: {docx_path}")

        except Exception as e:
            print(f"[ERROR] Failed to create DOCX: {e}")
            import traceback
            traceback.print_exc()

    def _create_docx_from_markdown(self, md_path: str) -> "Document":
        """
        使用纯 python-docx 从 markdown 创建 Word 文档。

        Args:
            md_path: Markdown 文件路径

        Returns:
            Document 对象
        """
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        import re

        doc = Document()

        # 读取 markdown 内容
        with open(md_path, "r", encoding="utf-8") as f:
            content = f.read()

        lines = content.split("\n")
        i = 0
        in_code_block = False
        code_content = []
        in_table = False
        table_rows = []

        while i < len(lines):
            line = lines[i]

            # 处理代码块
            if line.strip().startswith("```"):
                if in_code_block:
                    # 结束代码块
                    in_code_block = False
                    if code_content:
                        # 添加代码块内容
                        code_para = doc.add_paragraph()
                        code_para.style = "Normal"
                        code_run = code_para.add_run("\n".join(code_content))
                        code_run.font.name = "Courier New"
                        code_run.font.size = Pt(9)
                        code_content = []
                else:
                    # 开始代码块
                    in_code_block = True
                i += 1
                continue

            if in_code_block:
                code_content.append(line)
                i += 1
                continue

            # 处理表格
            if line.strip().startswith("|"):
                if not in_table:
                    in_table = True
                    table_rows = []
                # 跳过分隔行
                if re.match(r"^\|[\s\-:|]+\|$", line.strip()):
                    i += 1
                    continue
                # 解析表格行
                cells = [cell.strip() for cell in line.strip().split("|")[1:-1]]
                if cells:
                    table_rows.append(cells)
                i += 1
                continue
            elif in_table:
                # 表格结束，创建表格
                in_table = False
                if table_rows:
                    self._add_table_to_doc(doc, table_rows)
                    table_rows = []

            # 处理标题
            if line.startswith("# "):
                heading = doc.add_heading(line[2:].strip(), level=1)
                i += 1
                continue
            elif line.startswith("## "):
                heading = doc.add_heading(line[3:].strip(), level=2)
                i += 1
                continue
            elif line.startswith("### "):
                heading = doc.add_heading(line[4:].strip(), level=3)
                i += 1
                continue
            elif line.startswith("#### "):
                heading = doc.add_heading(line[5:].strip(), level=4)
                i += 1
                continue

            # 处理水平线
            if line.strip() == "---":
                # 添加水平线
                para = doc.add_paragraph()
                para.add_run("─" * 50)
                i += 1
                continue

            # 处理列表项
            if line.strip().startswith("- ") or line.strip().startswith("* "):
                text = line.strip()[2:]
                para = doc.add_paragraph(style="List Bullet")
                self._add_formatted_text(para, text)
                i += 1
                continue

            # 处理数字列表
            match = re.match(r"^\d+\.\s+(.+)$", line.strip())
            if match:
                para = doc.add_paragraph(style="List Number")
                self._add_formatted_text(para, match.group(1))
                i += 1
                continue

            # 处理普通段落
            if line.strip():
                para = doc.add_paragraph()
                self._add_formatted_text(para, line.strip())

            i += 1

        # 处理最后的表格
        if in_table and table_rows:
            self._add_table_to_doc(doc, table_rows)

        return doc

    def _add_table_to_doc(self, doc: "Document", rows: List[List[str]]) -> None:
        """
        向文档添加专业格式的表格。

        Args:
            doc: Document 对象
            rows: 表格行数据
        """
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.oxml.ns import nsdecls
        from docx.oxml import parse_xml
        import re

        if not rows or len(rows) < 1:
            return

        # 确定列数
        num_cols = max(len(row) for row in rows)
        if num_cols == 0:
            return

        # 创建表格
        table = doc.add_table(rows=len(rows), cols=num_cols)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 填充表格
        for i, row_data in enumerate(rows):
            row = table.rows[i]
            for j, cell_text in enumerate(row_data):
                if j < num_cols:
                    cell = row.cells[j]
                    # 清空默认段落
                    cell.text = ""
                    para = cell.paragraphs[0]

                    # 处理 markdown 加粗语法 **text**
                    self._add_formatted_cell_text(para, cell_text, Pt(10), is_header=(i == 0))

            # 设置表头行背景色
            if i == 0:
                for j in range(num_cols):
                    if j < len(row.cells):
                        cell = row.cells[j]
                        # 设置浅蓝色背景
                        shading_elm = parse_xml(
                            r'<w:shd {} w:fill="D9E2F3"/>'.format(nsdecls('w'))
                        )
                        cell._tc.get_or_add_tcPr().append(shading_elm)

        # 添加空行
        doc.add_paragraph()

    def _add_formatted_cell_text(self, para, text: str, font_size, is_header: bool = False) -> None:
        """
        向单元格段落添加格式化文本，处理 markdown 加粗语法。

        Args:
            para: 段落对象
            text: 文本内容
            font_size: 字体大小
            is_header: 是否为表头行
        """
        import re

        # 处理 markdown 加粗语法 **text**
        pattern = r'(\*\*(.+?)\*\*|([^*]+))'
        matches = re.findall(pattern, text)

        for match in matches:
            full, bold_text, normal_text = match
            if bold_text:
                run = para.add_run(bold_text)
                run.bold = True
                run.font.size = font_size
                run.font.name = "Arial"
            elif normal_text:
                run = para.add_run(normal_text)
                run.font.size = font_size
                run.font.name = "Arial"
                if is_header:
                    run.bold = True

    def _add_formatted_text(self, para, text: str) -> None:
        """
        向段落添加格式化文本，处理粗体和斜体。

        Args:
            para: 段落对象
            text: 文本内容
        """
        import re

        # 简单处理粗体 **text** 和斜体 *text*
        pattern = r"(\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`|([^*`]+))"
        matches = re.findall(pattern, text)

        for match in matches:
            full, bold, italic, code, normal = match
            if bold:
                run = para.add_run(bold)
                run.bold = True
            elif italic:
                run = para.add_run(italic)
                run.italic = True
            elif code:
                run = para.add_run(code)
                run.font.name = "Courier New"
            elif normal:
                para.add_run(normal)

    def _track_data_flow(self, messages: List) -> None:
        """Track data flow between agents"""
        # Create mapping from agent names to config keys
        agent_name_to_config_key = {
            "parser_agent": "parser",
            "analysis_agent": "analysis",
            "review_agent": "review",
            "report_agent": "report",
        }

        for msg in messages:
            if isinstance(msg.content, str) and msg.content.strip():
                # Try multiple ways to get message source
                source = getattr(msg, "source", None)

                # If source is empty, try agent_name
                if not source:
                    source = getattr(msg, "agent_name", None)

                # If still empty, try name
                if not source:
                    source = getattr(msg, "name", None)

                # If still empty, infer from content (via termination marker)
                if not source:
                    for agent_name, config in self.agent_configs.items():
                        if config.termination_marker in msg.content:
                            source = agent_name
                            break

                # Map agent name to config key
                config_key = agent_name_to_config_key.get(source, source)

                # Record output - skip 'user' source
                if source and source != "user" and config_key in self.agent_configs:
                    self.data_flow_manager.record_output(
                        config_key,
                        msg.content,
                        self.agent_configs[config_key].output_format,
                    )
                    print(f"[DEBUG] Recorded output from {source} (config_key={config_key})")

    def _extract_final_report(self, messages: List) -> str:
        """Extract and aggregate final report from all agent outputs"""
        # Collect outputs from all agents in order
        agent_outputs = {
            "parser_agent": None,
            "analysis_agent": None,
            "review_agent": None,
            "report_agent": None,
        }

        # Collect all outputs, keeping the latest from each agent
        for msg in messages:
            if isinstance(msg.content, str) and msg.content.strip():
                source = getattr(msg, "source", None)
                if not source:
                    source = getattr(msg, "agent_name", None)
                if not source:
                    source = getattr(msg, "name", None)

                # Store the latest output from each agent
                if source in agent_outputs:
                    # Remove termination markers
                    content = msg.content
                    for marker in ["[PARSER_DONE]", "[ANALYSIS_DONE]", "[REVIEW_DONE]", "[REPORT_DONE]"]:
                        content = content.replace(marker, "")
                    content = content.strip()

                    if content:
                        agent_outputs[source] = content

        # Get the latest outputs
        parser_output = agent_outputs.get("parser_agent")
        analysis_output = agent_outputs.get("analysis_agent")
        review_output = agent_outputs.get("review_agent")
        report_output = agent_outputs.get("report_agent")

        # Always use aggregation to ensure complete report structure
        print("[DEBUG] Aggregating outputs from all agents into final report")
        aggregated_report = self._aggregate_report(
            parser_output, analysis_output, review_output, report_output
        )

        return aggregated_report

    def _parse_text_to_table_data(self, text: str) -> Dict[str, str]:
        """
        从原始文本中提取键值对数据。

        Args:
            text: 原始文本内容

        Returns:
            提取的键值对字典
        """
        import re
        table_data = {}

        # 尝试多种格式的键值对提取
        # 格式1: "Key: Value" 或 "Key = Value"
        patterns = [
            r'^\s*([A-Za-z_\s]+):\s*(.+)$',  # Key: Value
            r'^\s*([A-Za-z_\s]+)\s*=\s*(.+)$',  # Key = Value
            r'^\s*-\s*([A-Za-z_\s]+):\s*(.+)$',  # - Key: Value (列表格式)
            r'^\s*\*\s*([A-Za-z_\s]+):\s*(.+)$',  # * Key: Value (markdown列表)
        ]

        lines = text.split('\n')
        for line in lines:
            line = line.strip()
            if not line:
                continue

            for pattern in patterns:
                match = re.match(pattern, line, re.IGNORECASE)
                if match:
                    key = match.group(1).strip().title()
                    value = match.group(2).strip()
                    # 清理值中的特殊字符
                    value = value.strip('`"\'')
                    if key and value:
                        table_data[key] = value
                    break

        # 如果没有提取到数据，尝试从 JSON 代码块中提取
        if not table_data:
            json_match = re.search(r'```(?:json)?\s*(\{[\s\S]*?\})\s*```', text)
            if json_match:
                try:
                    json_data = json.loads(json_match.group(1))
                    table_data = self._flatten_json(json_data)
                except (json.JSONDecodeError, ValueError):
                    pass

        return table_data

    def _flatten_json(self, data: dict, prefix: str = "") -> Dict[str, str]:
        """
        将嵌套的 JSON 数据扁平化为键值对。

        Args:
            data: JSON 数据
            prefix: 键前缀

        Returns:
            扁平化的键值对字典
        """
        result = {}
        for key, value in data.items():
            full_key = f"{prefix}{key}".replace('_', ' ').title() if prefix else key.replace('_', ' ').title()
            if isinstance(value, dict):
                result.update(self._flatten_json(value, f"{full_key} - "))
            elif isinstance(value, list):
                result[full_key] = ', '.join(str(v) for v in value)
            else:
                result[full_key] = str(value)
        return result

    def _aggregate_report(self, parser_output: str, analysis_output: str,
                         review_output: str, report_output: str) -> str:
        """
        Aggregate outputs from all agents into a professional, well-structured report.

        生成专业的 ITU 干扰分析报告，优化格式：
        - 使用表格展示结构化数据
        - 使用段落展示分析文本
        - 添加小标题编号（如 2.1, 2.2）
        - 总结内容加粗，细节描述正常字体
        - 添加 ITU 标准参考
        """
        import re
        report_parts = []

        # 添加报告标题和头部信息（不使用表格，更醒目）
        report_parts.append("# ITU Interference Analysis Report")
        report_parts.append("")
        report_parts.append("---")
        report_parts.append("")
        report_parts.append(f"**Report Date:** {self.current_date}")
        report_parts.append("")
        report_parts.append("**Organization:** Institute of Space Internet, Fudan University")
        report_parts.append("")
        report_parts.append("**Analysis System:** Calference - Satellite Interference Analysis Platform")
        report_parts.append("")
        report_parts.append(f"**Document ID:** ITU-RPT-{self.current_date.replace('-', '')}-{self.image_info.get('constellation', 'SAT').upper()[:3]}")
        report_parts.append("")
        report_parts.append("---")
        report_parts.append("")

        # Section 1: Executive Summary
        report_parts.append("## 1. Executive Summary")
        report_parts.append("")
        report_parts.append(f"**This report presents the interference analysis results for the {self.image_info.get('constellation', 'N/A').upper()} constellation's {self.image_info.get('terminal_type', 'system')}.** The analysis focuses on {self.image_info.get('analysis_type', 'interference metrics')} measurements collected over a 24-hour monitoring period.")
        report_parts.append("")
        report_parts.append("The assessment is conducted in accordance with ITU Radio Regulations and relevant ITU-R Recommendations to ensure compliance with international spectrum management requirements.")
        report_parts.append("")

        section_num = 2

        # Section 2: System Information
        report_parts.append(f"## {section_num}. System Under Test")
        report_parts.append("")
        report_parts.append(f"### {section_num}.1 Basic Information")
        report_parts.append("")
        report_parts.append("| Parameter | Value |")
        report_parts.append("|-----------|-------|")
        report_parts.append(f"| Constellation Name | {self.image_info.get('constellation', 'N/A').upper()} |")
        report_parts.append(f"| Terminal Type | {self.image_info.get('terminal_type', 'N/A')} |")
        report_parts.append(f"| Analysis Metric | {self.image_info.get('analysis_type', 'N/A')} |")
        report_parts.append(f"| Monitoring Period | 24 hours (00:00 - 23:59 UTC) |")
        report_parts.append(f"| Source Data File | {self.image_info.get('filename', 'N/A')} |")
        report_parts.append("")

        report_parts.append(f"### {section_num}.2 Applicable ITU Standards")
        report_parts.append("")
        report_parts.append("The following ITU-R Recommendations are referenced in this analysis:")
        report_parts.append("")
        report_parts.append("| Standard | Description |")
        report_parts.append("|----------|-------------|")
        report_parts.append("| ITU-R S.1503 | Functional description to be used in developing software tools for determining conformity of non-GSO FSS systems |")
        report_parts.append("| ITU-R S.1325 | Simulation methodologies for determining statistics of short-term interference |")
        report_parts.append("| ITU-R S.1528 | Satellite antenna radiation patterns for non-GSO orbit systems |")
        report_parts.append("| ITU-R SF.1395 | Minimum propagation attenuation due to atmospheric gases for frequency coordination |")
        report_parts.append("| ITU-R P.618 | Propagation data and prediction methods for earth-space telecommunication systems |")
        report_parts.append("")

        section_num += 1

        # Section 3: Data Analysis (from parser output)
        if parser_output:
            report_parts.append(f"## {section_num}. Measurement Data Analysis")
            report_parts.append("")
            report_parts.append(f"**This section presents the quantitative analysis of the {self.image_info.get('analysis_type', 'interference')} measurements extracted from the monitoring data.**")
            report_parts.append("")

            subsection = 1
            try:
                # 尝试解析 JSON
                parser_data = json.loads(parser_output)
                self._add_structured_data_section_v2(report_parts, parser_data, section_num, subsection)
            except (json.JSONDecodeError, KeyError, TypeError) as e:
                print(f"[DEBUG] JSON parsing failed: {e}, converting raw output to table format")
                # 尝试从原始文本中提取键值对
                table_data = self._parse_text_to_table_data(parser_output)
                if table_data:
                    self._add_extracted_data_table_v2(report_parts, table_data, section_num)
                else:
                    report_parts.append(parser_output)
                    report_parts.append("")

            section_num += 1

        # Section 4: Technical Analysis (from analysis agent)
        if analysis_output:
            analysis_clean = self._clean_agent_output(analysis_output, "analysis")
            if analysis_clean.strip():
                report_parts.append(f"## {section_num}. Technical Analysis")
                report_parts.append("")
                report_parts.append(f"**The following technical analysis evaluates the interference characteristics and potential impact on system performance.**")
                report_parts.append("")
                report_parts.append(analysis_clean)
                report_parts.append("")
                section_num += 1

        # Section 5: Compliance Assessment (from review agent)
        if review_output:
            review_clean = self._clean_agent_output(review_output, "review")
            if review_clean.strip():
                report_parts.append(f"## {section_num}. ITU Compliance Assessment")
                report_parts.append("")
                report_parts.append(f"**This section assesses the system's compliance with applicable ITU Radio Regulations and Recommendations.**")
                report_parts.append("")
                report_parts.append(review_clean)
                report_parts.append("")
                section_num += 1

        # Section 6: Conclusions and Recommendations (from report agent)
        if report_output:
            report_clean = self._clean_agent_output(report_output, "report")
            if report_clean.strip():
                report_parts.append(f"## {section_num}. Conclusions and Recommendations")
                report_parts.append("")
                report_parts.append(f"**Based on the comprehensive analysis presented above, the following conclusions and recommendations are provided.**")
                report_parts.append("")
                report_parts.append(report_clean)
                report_parts.append("")
                section_num += 1

        # Section: References
        report_parts.append("---")
        report_parts.append(f"## {section_num}. References")
        report_parts.append("")
        report_parts.append("### ITU-R Recommendations")
        report_parts.append("")
        report_parts.append("1. **ITU-R S.1503-3** (2013): Functional description to be used in developing software tools for determining conformity of non-geostationary-satellite orbit fixed-satellite service systems or networks with limits contained in Article 22 of the Radio Regulations")
        report_parts.append("")
        report_parts.append("2. **ITU-R S.1325-3** (2003): Simulation methodologies for determining statistics of short-term interference between co-frequency, codirectional non-geostationary-satellite orbit fixed-satellite service systems in circular orbits and other non-geostationary fixed-satellite service systems in circular orbits or geostationary-satellite orbit fixed-satellite service networks")
        report_parts.append("")
        report_parts.append("3. **ITU-R S.1528** (2001): Satellite antenna radiation patterns for non-geostationary orbit satellite antennas operating in the fixed-satellite service below 30 GHz")
        report_parts.append("")
        report_parts.append("4. **ITU-R SF.1395** (1999): Minimum propagation attenuation due to atmospheric gases for use in frequency coordination between the fixed-satellite service and the fixed service")
        report_parts.append("")
        report_parts.append("5. **ITU-R P.618-13** (2017): Propagation data and prediction methods required for the design of Earth-space telecommunication systems")
        report_parts.append("")
        section_num += 1

        # Appendix
        report_parts.append("---")
        report_parts.append(f"## {section_num}. Appendix: Report Metadata")
        report_parts.append("")
        report_parts.append("| Parameter | Value |")
        report_parts.append("|-----------|-------|")
        report_parts.append("| Report Type | Multi-Agent Dialogue-Based Analysis |")
        report_parts.append("| Analysis Pipeline | Parser → Analysis → Review → Report |")
        report_parts.append(f"| Generation Date | {self.current_date} |")
        report_parts.append("| Analysis System | Calference v1.0 |")
        report_parts.append("| Organization | Institute of Space Internet, Fudan University |")
        report_parts.append("")
        report_parts.append("---")
        report_parts.append("")
        report_parts.append("*This report was automatically generated by the Calference Satellite Interference Analysis Platform.*")
        report_parts.append("")
        report_parts.append("*© 2026 Institute of Space Internet, Fudan University. All rights reserved.*")
        report_parts.append("")

        # 合并所有部分
        final_report = "\n".join(report_parts)
        # 清理多余空行
        while "\n\n\n" in final_report:
            final_report = final_report.replace("\n\n\n", "\n\n")

        return final_report

    def _add_structured_data_section_v2(self, report_parts: List[str], parser_data: dict,
                                        section_num: int, subsection: int) -> None:
        """添加结构化数据章节（从 JSON 解析），带小标题编号"""
        # Numerical Data Table
        if "numerical_data" in parser_data:
            report_parts.append(f"### {section_num}.{subsection} Numerical Measurements")
            report_parts.append("")
            report_parts.append("**Summary:** The table below presents the key statistical parameters extracted from the monitoring data.")
            report_parts.append("")
            num_data = parser_data["numerical_data"]
            report_parts.append("| Parameter | Value | Unit | Description |")
            report_parts.append("|-----------|-------|------|-------------|")
            report_parts.append(f"| Metric Type | {num_data.get('metric_name', 'N/A')} | - | Primary interference metric |")
            report_parts.append(f"| Minimum | {num_data.get('min_value', 'N/A')} | {num_data.get('unit', 'dB')} | Lowest recorded value |")
            report_parts.append(f"| Maximum | {num_data.get('max_value', 'N/A')} | {num_data.get('unit', 'dB')} | Highest recorded value |")
            report_parts.append(f"| Average | {num_data.get('avg_value', 'N/A')} | {num_data.get('unit', 'dB')} | Mean value over monitoring period |")
            report_parts.append(f"| Range | {num_data.get('range', 'N/A')} | {num_data.get('unit', 'dB')} | Difference between max and min |")
            report_parts.append("")
            subsection += 1

        # Temporal Characteristics
        if "temporal_characteristics" in parser_data:
            report_parts.append(f"### {section_num}.{subsection} Temporal Characteristics")
            report_parts.append("")
            report_parts.append("**Summary:** The temporal analysis identifies patterns and anomalies in the time-series data.")
            report_parts.append("")
            temporal = parser_data["temporal_characteristics"]
            abnormal = ', '.join(temporal.get('abnormal_periods', [])) if temporal.get("abnormal_periods") else "None identified"
            report_parts.append("| Characteristic | Value | Interpretation |")
            report_parts.append("|----------------|-------|----------------|")
            report_parts.append(f"| Abnormal Periods | {abnormal} | Time periods with significant deviations |")
            report_parts.append(f"| Duration Pattern | {temporal.get('duration_type', 'N/A').capitalize()} | Nature of interference events |")
            report_parts.append(f"| Magnitude Level | {temporal.get('magnitude', 'N/A').capitalize()} | Severity of observed anomalies |")
            report_parts.append("")
            subsection += 1

        # Visual Patterns
        if "visual_patterns" in parser_data and parser_data["visual_patterns"]:
            report_parts.append(f"### {section_num}.{subsection} Observed Visual Patterns")
            report_parts.append("")
            report_parts.append("**Summary:** The following patterns were identified through visual inspection of the monitoring data:")
            report_parts.append("")
            patterns = parser_data["visual_patterns"]
            if isinstance(patterns, list):
                for i, pattern in enumerate(patterns, 1):
                    report_parts.append(f"{i}. {pattern}")
            else:
                report_parts.append(f"- {str(patterns)}")
            report_parts.append("")

    def _add_extracted_data_table_v2(self, report_parts: List[str], table_data: dict,
                                     section_num: int) -> None:
        """添加从文本提取的数据表格，按类别分组，带小标题编号"""
        # 按前缀分组数据
        groups = {}
        for key, value in table_data.items():
            if " - " in key:
                prefix, suffix = key.split(" - ", 1)
                if prefix not in groups:
                    groups[prefix] = {}
                groups[prefix][suffix] = value
            else:
                if "General" not in groups:
                    groups["General"] = {}
                groups["General"][key] = value

        # 输出分组表格
        subsection = 1
        for group_name, group_data in groups.items():
            if not group_data:
                continue
            report_parts.append(f"### {section_num}.{subsection} {group_name}")
            report_parts.append("")
            report_parts.append(f"**Summary:** Key parameters extracted for {group_name.lower()} analysis.")
            report_parts.append("")
            report_parts.append("| Parameter | Value |")
            report_parts.append("|-----------|-------|")
            for key, value in group_data.items():
                # 清理值中的字典/列表表示
                if isinstance(value, str) and value.startswith("{"):
                    value = "See detailed analysis"
                report_parts.append(f"| {key} | {value} |")
            report_parts.append("")
            subsection += 1

    def _clean_agent_output(self, output: str, agent_type: str) -> str:
        """
        清理 agent 输出，移除重复内容和格式问题。

        Args:
            output: 原始输出
            agent_type: agent 类型 (analysis/review/report)

        Returns:
            清理后的输出
        """
        import re

        if not output:
            return ""

        cleaned = output.strip()

        # 移除所有代码块标记（包括 ```json, ```markdown 等）
        cleaned = re.sub(r'```(?:json|markdown|text)?\s*', '', cleaned)

        # 移除独立的 JSON 对象（不在代码块中的）
        cleaned = re.sub(r'^\s*\{[\s\S]*?\}\s*$', '', cleaned, flags=re.MULTILINE)
        # 移除多行 JSON 对象
        cleaned = re.sub(r'\{\s*"[^"]+"\s*:\s*\{[\s\S]*?\}\s*\}', '', cleaned)
        cleaned = re.sub(r'\{\s*"basic_info"[\s\S]*?\}\s*\}', '', cleaned)
        # 移除残留的单独大括号
        cleaned = re.sub(r'^\s*[\{\}]\s*$', '', cleaned, flags=re.MULTILINE)

        # 移除重复的章节标题
        headers_to_remove = [
            r'^#+\s*ITU\s*Interference Analysis Report\s*$',
            r'^#+\s*Interference Analysis Report\s*$',
            r'^#+\s*\d+\.\s*Basic Information\s*$',
            r'^#+\s*\d+\.\s*Data Analysis\s*$',
            r'^#+\s*\d+\.\s*Compliance Assessment\s*$',
            r'^#+\s*\d+\.\s*Conclusions.*$',
            r'^#+\s*\d+\.\s*Appendix.*$',
            r'^#+\s*Data Analysis Results\s*$',
            r'^#+\s*Compliance Review\s*$',
            r'^#+\s*Report Overview\s*$',
            r'^#+\s*Data Extraction Results\s*$',
            r'^#+\s*Technical Analysis\s*$',
            r'^#+\s*ITU Compliance Assessment\s*$',
        ]
        for pattern in headers_to_remove:
            cleaned = re.sub(pattern, '', cleaned, flags=re.MULTILINE | re.IGNORECASE)

        # 移除只有表头没有内容的表格
        cleaned = re.sub(r'\|\s*Field\s*\|\s*Value\s*\|\s*\n\|[-\s|]+\|\s*\n(?!\|)', '', cleaned)
        cleaned = re.sub(r'\|\s*Item\s*\|\s*Value\s*\|\s*\n\|[-\s|]+\|\s*\n(?!\|)', '', cleaned)
        cleaned = re.sub(r'\|\s*Item\s*\|\s*Description\s*\|\s*\n\|[-\s|]+\|\s*\n(?!\|)', '', cleaned)

        # 移除 "End with:" 等提示文本
        cleaned = re.sub(r'^End with:\s*$', '', cleaned, flags=re.MULTILINE)

        # 移除占位符文本
        placeholder_patterns = [
            r'\[1 paragraph explaining.*?\]',
            r'\[Visual cue \d+ with specifics\]',
            r'\[.*?placeholder.*?\]',
            r'<\|observation\|>',
            r'<\|.*?\|>',
        ]
        for pattern in placeholder_patterns:
            cleaned = re.sub(pattern, '', cleaned, flags=re.IGNORECASE)

        # 处理行
        lines = cleaned.split('\n')
        result_lines = []
        list_items = []
        in_list = False

        for line in lines:
            stripped = line.strip()

            # 跳过空行（在开头）
            if not result_lines and not stripped:
                continue

            # 检测列表项
            is_list_item = stripped.startswith('- ') or stripped.startswith('* ') or re.match(r'^\d+\.\s', stripped)

            if is_list_item:
                in_list = True
                # 提取列表内容
                if stripped.startswith('- ') or stripped.startswith('* '):
                    content = stripped[2:]
                else:
                    content = re.sub(r'^\d+\.\s*', '', stripped)
                list_items.append(content)
            else:
                # 如果之前在列表中，现在不是列表项了
                if in_list and list_items:
                    # 如果列表项包含 "Key: Value" 格式，转换为表格
                    if all(':' in item for item in list_items[:3]) and len(list_items) >= 2:
                        result_lines.append("")
                        result_lines.append("| Item | Description |")
                        result_lines.append("|------|-------------|")
                        for item in list_items:
                            if ':' in item:
                                key, val = item.split(':', 1)
                                result_lines.append(f"| {key.strip()} | {val.strip()} |")
                            else:
                                result_lines.append(f"| - | {item} |")
                        result_lines.append("")
                    else:
                        # 保持为列表，但使用数字编号
                        result_lines.append("")
                        for i, item in enumerate(list_items, 1):
                            result_lines.append(f"{i}. {item}")
                        result_lines.append("")
                    list_items = []
                    in_list = False

                if stripped:
                    result_lines.append(line)
                elif result_lines and result_lines[-1].strip():
                    result_lines.append("")

        # 处理最后的列表
        if list_items:
            if all(':' in item for item in list_items[:3]) and len(list_items) >= 2:
                result_lines.append("")
                result_lines.append("| Item | Description |")
                result_lines.append("|------|-------------|")
                for item in list_items:
                    if ':' in item:
                        key, val = item.split(':', 1)
                        result_lines.append(f"| {key.strip()} | {val.strip()} |")
                    else:
                        result_lines.append(f"| - | {item} |")
                result_lines.append("")
            else:
                result_lines.append("")
                for i, item in enumerate(list_items, 1):
                    result_lines.append(f"{i}. {item}")
                result_lines.append("")

        cleaned = '\n'.join(result_lines)

        # 清理多余空行
        while '\n\n\n' in cleaned:
            cleaned = cleaned.replace('\n\n\n', '\n\n')

        return cleaned.strip()


async def main() -> None:
    """Main entry point"""
    configure_proxies()

    print("==== ITU Interference Report Generation (Configurable Dialogue-Based Multi-Agent) ====")

    image_path = os.path.join(INPUT_IMAGE_DIR, "oneweb_total_earth_cinr.png")
    filename = os.path.basename(image_path)
    image_info = parse_image_info(filename)
    current_date = datetime.now().strftime("%Y-%m-%d")

    print(f"[INFO] Config: model={LLM_MODEL_NAME} | compress_image=False | configurable_mode=True")

    # Use configurable pipeline
    pipeline = ConfigurableDialoguePipeline(image_info, current_date, image_path)
    await pipeline.run(
        image_path=image_path,
        compress_img=False,
    )


if __name__ == "__main__":
    asyncio.run(main())
