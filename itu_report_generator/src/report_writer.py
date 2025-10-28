# Copyright (C) 2025 FDU-INC
#
# This program is free software: you can redistribute it and/or modify
# it under the terms of the GNU General Public License as published by
# the Free Software Foundation, either version 3 of the License, or
# (at your option) any later version.
#
# This program is distributed in the hope that it will be useful,
# but WITHOUT ANY WARRANTY; without even the implied warranty of
# MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
# GNU General Public License for more details.
#
# You should have received a copy of the GNU General Public License
# along with this program.  If not, see <https://www.gnu.org/licenses/>.

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re

class WordReportWriter:
    def generate_report(self, image_path, report_content, output_dir, lang='zh'):
        """生成并保存Word报告（支持markdown风格格式化）"""
        base_filename = os.path.basename(image_path)
        report_filename = base_filename.replace('.png', '.docx')
        output_path = os.path.join(output_dir, report_filename)

        doc = Document()
        if lang == 'en':
            doc.add_heading(f'Interference Analysis Report: {base_filename}', level=1)
            doc.add_heading('Analysis Chart', level=2)
        else:
            doc.add_heading(f'干扰分析报告: {base_filename}', level=1)
            doc.add_heading('分析图表', level=2)
        doc.add_picture(image_path, width=Inches(6.0))
        self._add_markdown_content(doc, report_content)
        doc.save(output_path)
        print(f"报告已生成并保存至: {output_path}")
        return output_path

    def generate_report_en(self, image_path, report_content, output_dir):
        # 清理markdown代码块包裹和末尾TERMINATE
        content = report_content.strip()
        if content.startswith('```markdown'):
            content = content[len('```markdown'):].lstrip('\n')
        if content.startswith('```'):
            content = content[len('```'):].lstrip('\n')
        if content.endswith('```'):
            content = content[:-3].rstrip('\n')
        if content.rstrip().endswith('TERMINATE'):
            content = content.rstrip()
            content = content[:content.rfind('TERMINATE')].rstrip('\n')
        base_filename = os.path.basename(image_path)
        report_filename = base_filename.replace('.png', '.docx')
        output_path = os.path.join(output_dir, report_filename)
        doc = Document()
        lines = content.split('\n')
        # 找到3. Simulation Results Analysis标题行
        insert_idx = None
        for idx, line in enumerate(lines):
            if line.strip().startswith('## 3. Simulation Results Analysis'):
                insert_idx = idx
                break
        if insert_idx is not None:
            # 1. 先插入3章前的内容
            before_content = '\n'.join(lines[:insert_idx])
            if before_content.strip():
                self._add_markdown_content_en(doc, before_content)
            # 2. 插入3章标题和图片
            doc.add_heading(lines[insert_idx].strip()[3:], level=2)
            doc.add_paragraph()  # 空一行
            pic = doc.add_picture(image_path, width=Inches(4.0))
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            doc.add_paragraph()  # 图片后再空一行
            # 3. 插入3章及后续内容
            rest_content = '\n'.join(lines[insert_idx+1:])
            self._add_markdown_content_en(doc, rest_content)
        else:
            # 没找到标题，全部插入
            self._add_markdown_content_en(doc, content)
        # 调整所有标题字号和颜色
        heading_font_sizes = {
            'Heading 1': 22,
            'Heading 2': 18,
            'Heading 3': 16,
            'Heading 4': 14,
        }
        for para in doc.paragraphs:
            style_name = para.style.name
            if style_name in heading_font_sizes:
                for run in para.runs:
                    run.font.size = Pt(heading_font_sizes[style_name])
                    if style_name == 'Heading 1':
                        run.font.color.rgb = RGBColor(31, 78, 121)
            if para.text.strip().startswith("Organization:") or para.text.strip().startswith("Report Date:"):
                for run in para.runs:
                    run.font.size = Pt(10)
        doc.save(output_path)
        print(f"报告已生成并保存至: {output_path}")
        return output_path

    def _add_markdown_content(self, doc, content):
        """将markdown风格内容格式化插入docx文档"""
        lines = content.split('\n')
        table_mode = False
        table_rows = []
        for i, line in enumerate(lines):
            # 标题处理
            if re.match(r'^# ', line):
                doc.add_heading(line[2:].strip(), level=1)
            elif re.match(r'^## ', line):
                doc.add_heading(line[3:].strip(), level=2)
            elif re.match(r'^### ', line):
                doc.add_heading(line[4:].strip(), level=3)
            elif re.match(r'^#### ', line):
                doc.add_heading(line[5:].strip(), level=4)
            # 表格处理
            elif re.match(r'^\|', line):
                table_mode = True
                table_rows.append([cell.strip() for cell in line.strip('|').split('|')])
            elif table_mode and (line.strip() == '' or not line.startswith('|')):
                # 表格结束，插入表格
                if len(table_rows) > 1:
                    table = doc.add_table(rows=1, cols=len(table_rows[0]))
                    table.style = 'Table Grid'
                    hdr_cells = table.rows[0].cells
                    for j, cell in enumerate(table_rows[0]):
                        hdr_cells[j].text = cell
                        run = hdr_cells[j].paragraphs[0].runs[0]
                        run.bold = True
                    for row in table_rows[1:]:
                        row_cells = table.add_row().cells
                        for j, cell in enumerate(row):
                            row_cells[j].text = cell
                table_mode = False
                table_rows = []
                if line.strip() != '':
                    doc.add_paragraph(line)
            elif table_mode:
                table_rows.append([cell.strip() for cell in line.strip('|').split('|')])
            # 列表项加粗
            elif re.match(r'^- ', line):
                p = doc.add_paragraph()
                run = p.add_run(line)
                run.bold = True
            # 普通正文
            elif line.strip() != '':
                # 检查是否有**加粗**语法
                parts = re.split(r'(\*\*[^*]+\*\*)', line)
                p = doc.add_paragraph()
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = p.add_run(part[2:-2])
                        run.bold = True
                    else:
                        p.add_run(part)

    def _add_markdown_content_en(self, doc, content):
        """将英文 markdown 风格内容格式化插入 docx 文档（优化标题、表格、列表、粗体、空行分段，表格内容支持粗体，自动补齐单元格）"""
        lines = content.split('\n')
        table_mode = False
        table_rows = []
        para_buffer = []
        for i, line in enumerate(lines):
            line_strip = line.strip()
            # 标题处理
            if re.match(r'^# ', line_strip):
                self._flush_paragraph_buffer(doc, para_buffer)
                doc.add_heading(line_strip[2:].strip(), level=1)
            elif re.match(r'^## ', line_strip):
                self._flush_paragraph_buffer(doc, para_buffer)
                doc.add_heading(line_strip[3:].strip(), level=2)
            elif re.match(r'^### ', line_strip):
                self._flush_paragraph_buffer(doc, para_buffer)
                doc.add_heading(line_strip[4:].strip(), level=3)
            elif re.match(r'^#### ', line_strip):
                self._flush_paragraph_buffer(doc, para_buffer)
                doc.add_heading(line_strip[5:].strip(), level=4)
            elif re.match(r'^##### ', line_strip):
                self._flush_paragraph_buffer(doc, para_buffer)
                doc.add_heading(line_strip[6:].strip(), level=4)
            # 表格处理
            elif re.match(r'^\|', line_strip):
                # 跳过 markdown 表格分隔线 | --- | --- | ... |
                cells = [cell.strip() for cell in line_strip.strip('|').split('|')]
                if all(cell == '---' or cell == '' for cell in cells):
                    continue  # 跳过只包含 --- 的分隔线
                self._flush_paragraph_buffer(doc, para_buffer)
                table_mode = True
                table_rows.append(cells)
            elif table_mode and (line_strip == '' or not line_strip.startswith('|')):
                # 表格结束，插入表格
                if len(table_rows) > 1:
                    ncols = len(table_rows[0])
                    table = doc.add_table(rows=1, cols=ncols)
                    table.style = 'Table Grid'
                    hdr_cells = table.rows[0].cells
                    for j, cell in enumerate(table_rows[0]):
                        hdr_cells[j].text = cell
                        run = hdr_cells[j].paragraphs[0].runs[0]
                        run.bold = True
                    for row in table_rows[1:]:
                        row_cells = table.add_row().cells
                        # 补齐或截断每行单元格
                        for j in range(ncols):
                            cell = row[j] if j < len(row) else ""
                            if cell.startswith('**') and cell.endswith('**'):
                                run = row_cells[j].paragraphs[0].add_run(cell[2:-2])
                                run.bold = True
                            else:
                                row_cells[j].text = cell
                table_mode = False
                table_rows = []
                if line_strip != '':
                    para_buffer.append(line_strip)
            elif table_mode:
                table_rows.append([cell.strip() for cell in line_strip.strip('|').split('|')])
            # 列表项
            elif re.match(r'^[-*] ', line_strip):
                self._flush_paragraph_buffer(doc, para_buffer)
                p = doc.add_paragraph(style='List Bullet')
                content = line_strip[2:].strip()
                # 检查是否是粗体开头，支持 - **xxx**: yyy
                m = re.match(r'^\*\*([^*]+)\*\*(.*)', content)
                if m:
                    bold_text = m.group(1)
                    rest = m.group(2)
                    # 如果后面是冒号+正文，分开加粗
                    if rest.strip().startswith(':'):
                        run = p.add_run(bold_text)
                        run.bold = True
                        p.add_run(':')
                        after_colon = rest.strip()[1:].lstrip()
                        if after_colon:
                            p.add_run(after_colon)
                    else:
                        run = p.add_run(bold_text)
                        run.bold = True
                        if rest:
                            p.add_run(rest)
                else:
                    p.add_run(content)
            # 空行，分段
            elif line_strip == '':
                self._flush_paragraph_buffer(doc, para_buffer)
            # 水平分割线 ---
            elif line_strip == '---':
                self._flush_paragraph_buffer(doc, para_buffer)
                doc.add_paragraph()  # 插入一个空段落，不插入“---”文本
            # 普通正文
            else:
                para_buffer.append(line_strip)
        # 处理结尾残留
        self._flush_paragraph_buffer(doc, para_buffer)

    def _flush_paragraph_buffer(self, doc, para_buffer):
        if para_buffer:
            text = ' '.join(para_buffer)
            # 处理粗体
            parts = re.split(r'(\*\*[^*]+\*\*)', text)
            p = doc.add_paragraph()
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    p.add_run(part)
            para_buffer.clear()
