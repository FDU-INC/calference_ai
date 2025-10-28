# Copyright (C) 2025 FDU-INC
#
# This program is free software: you can redistribute it and/or modify
# it under the terms of the GNU General Public License as published by
# the Free Software Foundation, either version 3 of the License, or
# (at your option) any later version.
#
# This program is distributed in the hope that it will be useful,
# but WITHOUT ANY WARRANTY; without even the implied warranty of
# MERCHANTABILITY or FITNESS FOR A PARTICULAR PURPOSE.  See the
# GNU General Public License for more details.
#
# You should have received a copy of the GNU General Public License
# along with this program.  If not, see <https://www.gnu.org/licenses/>.


from . import utils
from . import prompt_builder
from .llm_inference import LLMInference
from .report_writer import WordReportWriter
import os
from docx import Document
from docx.shared import Inches

class InterferenceReportOrchestrator:
    def __init__(self, model_path, organization_name):
        self.llm = LLMInference(model_path)
        self.writer = WordReportWriter()
        self.organization_name = organization_name

    def generate_single_report(self, image_path, output_dir, lang='zh'):
        """为单个图片文件生成完整的分析报告"""
        try:
            # 1. 解析图片信息
            if lang == 'en':
                image_info = utils.parse_image_path(image_path, lang='en')
            else:
                image_info = utils.parse_image_path(image_path)
            system_name = image_info['constellation']
            other_system_name = "OneWeb" if system_name == "Omni" else "Omni"
            
            # 2. 动态生成图表清单
            image_types = ['cinr', 'cir', 'cnr', 'epfd', 'inr', 'link_count', 'pfd', 'temp']
            terminals = ['earth', 'satellite']
            image_list = [f"- {system_name.lower()}_total_{term}_{img_type}.png" for term in terminals for img_type in image_types]
            image_list_str = "\n".join(image_list)

            # 3. 构建Prompt
            if lang == 'en':
                prompt = prompt_builder.get_analysis_prompt_en(
                    system_name=system_name,
                    other_system_name=other_system_name,
                    report_date=utils.get_current_report_date(),
                    organization=self.organization_name,
                    image_info=image_info,
                    image_list_str=image_list_str
                )
                prompt += "\n\nPlease write the entire report in English only."
            else:
                prompt = prompt_builder.get_analysis_prompt(
                    system_name=system_name,
                    other_system_name=other_system_name,
                    report_date=utils.get_current_report_date(),
                    organization=self.organization_name,
                    image_info=image_info,
                    image_list_str=image_list_str
                )
            # 4. 调用LLM生成分析内容
            report_content = self.llm.generate_analysis(image_path, prompt)
            print(report_content)
            # 5. 写入Word文档
            if lang == 'en':
                self.writer.generate_report_en(image_path, report_content, output_dir)
            else:
                self.writer.generate_report(image_path, report_content, output_dir, lang=lang)

        except Exception as e:
            print(f"处理文件 {image_path} 时发生错误: {e}")

    def generate_multi_image_report(self, image_dir, output_dir, include_keywords=None, max_images=8, lang='zh'):
        """批量分析多张图片，生成一个Word报告，报告标题与原来一致，正文为结构化多图分析"""
        if include_keywords is None:
            include_keywords = []
        all_image_files = [f for f in os.listdir(image_dir) if f.endswith('.png')]
        filtered_files = [f for f in all_image_files if all(kw.lower() in f.lower() for kw in include_keywords)]
        filtered_files = filtered_files[:max_images]
        if not filtered_files:
            print(f"未找到包含关键词{include_keywords}的图片文件。")
            return
        print(f"本次将分析{len(filtered_files)}张图片: {filtered_files}")

        # 收集所有图片信息
        image_infos = []
        for image_file in filtered_files:
            image_path = os.path.join(image_dir, image_file)
            image_info = utils.parse_image_path(image_path)
            image_infos.append(image_info)
        # 假设所有图片属于同一星座
        system_name = image_infos[0]['constellation']
        other_system_name = "OneWeb" if system_name == "Omni" else "Omni"
        # 动态生成图表清单
        image_types = ['cinr', 'cir', 'cnr', 'epfd', 'inr', 'link_count', 'pfd', 'temp']
        terminals = ['earth', 'satellite']
        image_list = [f"- {system_name.lower()}_total_{term}_{img_type}.png" for term in terminals for img_type in image_types]
        image_list_str = "\n".join(image_list)

        # 构建多图分析文本 prompt
        if lang == 'en':
            text_prompt = prompt_builder.get_multi_image_analysis_prompt_en(
                system_name=system_name,
                other_system_name=other_system_name,
                report_date=utils.get_current_report_date(),
                organization=self.organization_name,
                image_infos=image_infos,
                image_list_str=image_list_str
            )
            text_prompt += "\n\nPlease write the entire report in English only."
        else:
            text_prompt = prompt_builder.get_multi_image_analysis_prompt(
                system_name=system_name,
                other_system_name=other_system_name,
                report_date=utils.get_current_report_date(),
                organization=self.organization_name,
                image_infos=image_infos,
                image_list_str=image_list_str
            )
        # 构建Qwen2.5 VL多模态messages
        image_paths = [os.path.join(image_dir, f) for f in filtered_files]
        messages = prompt_builder.build_qwen_vl_multimodal_messages(image_paths, text_prompt)
        # LLM生成整体分析内容
        report_content = self.llm.generate_analysis(messages, None)
        # # 生成Word报告
        # base_title = f"干扰分析报告: {', '.join(filtered_files)}"
        doc = Document()
        # doc.add_heading(base_title, level=1)
        # doc.add_heading('分析图表', level=2)

        # 处理report_content，将图片插入到"3. 仿真结果分析"节后
        section_title = '3. 仿真结果分析'
        # 支持不同markdown风格的标题
        if section_title in report_content:
            before, after = report_content.split(section_title, 1)
            # 插入前半部分
            if before.strip():
                self.writer._add_markdown_content(doc, before)
            # 插入3. 仿真结果分析标题
            doc.add_heading(section_title, level=1)
            # 插入所有图片
            for image_file in filtered_files:
                image_path = os.path.join(image_dir, image_file)
                doc.add_paragraph(image_file)
                doc.add_picture(image_path, width=Inches(6.0))
            # 插入后半部分
            self.writer._add_markdown_content(doc, after)
        else:
            # fallback: 全部插入
            self.writer._add_markdown_content(doc, report_content)

        # 保存Word文档
        report_filename = "multi_image_report.docx"
        output_path = os.path.join(output_dir, report_filename)
        doc.save(output_path)
        print(f"多图分析报告已生成并保存至: {output_path}")
        return output_path

